import sys
sys.path.insert(0, r'G:\DASTUR\URIT 50')
try:
    from docx import Document
    for fpath in [
        r'G:\Қилинган анализлар\05.03.2026\260304001324 Ismoilova Iroda.docx',
        r'G:\Қилинган анализлар\04.03.2026\260301001202 Test qilish uchun.docx'
    ]:
        print(f"\n{'='*60}")
        print(f"FILE: {fpath}")
        print('='*60)
        try:
            doc = Document(fpath)
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    print(f"  P{i}: {para.text[:100]}")
            print(f"\n  TABLES: {len(doc.tables)}")
            for ti, table in enumerate(doc.tables):
                print(f"\n  Table {ti+1} ({len(table.rows)} rows x {len(table.columns)} cols):")
                for ri, row in enumerate(table.rows[:5]):
                    cells = [c.text[:30] for c in row.cells]
                    print(f"    Row{ri}: {cells}")
                if len(table.rows) > 5:
                    print(f"    ... ({len(table.rows)-5} more rows)")
        except Exception as e:
            print(f"  ERROR: {e}")
except ImportError:
    print("python-docx not available, trying another approach")
