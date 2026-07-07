import os
import re
from datetime import datetime
import threading
import queue
import sys
import time

import serial
from docx import Document
from docx.shared import RGBColor, Pt
import mysql.connector
from monoblok_db_config import DB_CONFIG



# --- Fayl joylashgan papka ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ====== SOZLAMALAR (Tizim Sozlamalari oynasi orqali, analizator_config.json) ======
try:
    from monoblok_db_config import get_analyzer
    _siydik_cfg = get_analyzer("siydik")
except Exception as _e:
    print(f"[OGOHLANTIRISH] siydik config yuklanmadi: {_e}")
    _siydik_cfg = {}

COM_PORT = _siydik_cfg.get("com_port", "COM4")   # mijozda farq qilishi mumkin
BAUDRATE = int(_siydik_cfg.get("baudrate", 9600))

# Asosiy papkalar (frozen-aware — mijozda G: bo'lmasligi mumkin)
BASE_DIR = os.path.join(os.environ.get("ProgramData", r"C:\ProgramData"), "AzizMedLine", "URIT_natijalar")
os.makedirs(BASE_DIR, exist_ok=True)
RAW_LOG_DIR = r"G:\DASTUR\URIT 50\Urit\RAW_LOGS"

# Shablonni **shu skript turgan papkadan** olamiz (URIT 50 papkasidagi blank.docx)
TEMPLATE_PATH = os.path.join(SCRIPT_DIR, "blank.docx")

# ====== DATABASE CONFIG ======
# monoblok_db_config.py dan import qilinadi

# Global: skaner orqali kelgan barcode
barcode_queue = queue.Queue()
current_barcode = None

def replace_in_runs(paragraph, replacements: dict):
    """
    Paragraph ichidagi runlarni buzmasdan (format saqlanib),
    faqat matn qismini almashtiradi.
    """
    for run in paragraph.runs:
        for key, val in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, str(val))


# ====== RAW saqlash ======
def save_raw_block(block: str) -> str:
    dt = datetime.now()
    month_folder = os.path.join(RAW_LOG_DIR, dt.strftime("%Y%m"))
    os.makedirs(month_folder, exist_ok=True)
    fname = "urit_raw_" + dt.strftime("%Y%m%d_%H%M%S") + ".txt"
    full_path = os.path.join(month_folder, fname)
    with open(full_path, "w", encoding="utf-8") as f:
        f.write(block)
    return full_path


# ====== Sana bo'yicha papka ======
def get_save_dir() -> str:
    """
    Kunlik saqlash papkasini olish.
    Avtomatik: G:\\Қилинган анализлар\\Siydik tahlillari\\DD.MM.YYYY
    """
    # Kunlik sana bilan papka
    today = datetime.now().strftime("%d.%m.%Y")
    save_dir = os.path.join(BASE_DIR, today)
    
    # Papkani yaratish
    try:
        os.makedirs(save_dir, exist_ok=True)
        return save_dir
    except Exception as e:
        print(f"❌ Xatolik: Papka yaratishda muammo: {e}")
        # Agar xato bo'lsa, standart papkadan foydalanamiz
        save_dir = BASE_DIR
        os.makedirs(save_dir, exist_ok=True)
        return save_dir


# ====== DOC ichida matn almashtirish (FORMAT SAQLANIB) ======
def replace_text_in_doc(doc: Document, old: str, new: str):
    """
    Placeholderlarni almashtirish, formatni saqlab qolish.
    Faqat placeholder matnini almashtiradi, qolgan format o'zgarmaydi.
    """
    def _replace_in_paragraph(para):
        full_text = para.text
        
        if old not in full_text:
            return
        
        # Har bir run ichida almashtirish
        for run in para.runs:
            if old in run.text:
                # Formatni saqlab qolish
                original_font = run.font.name
                original_size = run.font.size
                original_bold = run.bold
                original_color = run.font.color.rgb if run.font.color.rgb else None
                
                # Matnni almashtirish
                run.text = run.text.replace(old, new)
                
                # Formatni qayta tiklash
                if original_font:
                    run.font.name = original_font
                if original_size:
                    run.font.size = original_size
                if original_bold is not None:
                    run.bold = original_bold
                if original_color:
                    run.font.color.rgb = original_color

    # Body - faqat placeholderlar bo'lgan paragraphlarni almashtirish
    for p in doc.paragraphs:
        if old in p.text:
            _replace_in_paragraph(p)

    # Jadval ichida
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        _replace_in_paragraph(p)

    # Header / Footer - barcha joylarda qidirish
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                if old in p.text:
                    _replace_in_paragraph(p)
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if old in p.text:
                                _replace_in_paragraph(p)
    
    # Barcha jadvallarda ham qidirish (vaqt uchun)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        _replace_in_paragraph(p)


# ====== URIT MA'LUMOTINI PARSE QILISH ======
ANALYTES = [
    "ACR",  # avval uzun kod
    "LEU", "KET", "NIT", "URO", "BIL",
    "PRO", "GLU", "SG", "pH", "BLD",
    "Vc", "MA", "Ca", "CR",
]

# 11-parametrli poloskalar: faqat asosiy ko'rsatkichlar
ANALYTES_11 = ["LEU", "KET", "NIT", "URO", "BIL", "PRO", "GLU", "SG", "pH", "BLD", "Vc"]


def detect_strip_type(data: dict) -> int:
    """Ca va MA qiymatlariga qarab poloskalar turini aniqlash.
    14-param: Ca '<' / '>' yoki raqamli diapazon (masalan '<=1.0 mmol/L').
    11-param: Ca oddiy son ('2.5 mmol/L') yoki bo'sh.
    """
    ca_val = str(data.get("Ca", "")).strip()
    if ca_val and ("<" in ca_val or ">" in ca_val or re.search(r"\d+-\d+", ca_val)):
        return 14
    ma_val = str(data.get("MA", "")).strip()
    if ma_val and (">" in ma_val or re.search(r"\d+-\d+", ma_val)):
        return 14
    return 11


# ====== MIKROSKOPIYA MAPPING (URIT jadvalidan) ======
# LEU va BLD natijalariga asosan mikroskopiya qiymatlarini aniqlash

def get_microscopy_values(leu_value: str, bld_value: str) -> dict:
    """
    URIT-50 dan kelgan LEU va BLD qiymatlariga asosan 
    mikroskopiya qiymatlarini hisoblash.
    
    Args:
        leu_value: Masalan "0 CELL/µL", "+1     70 CELL/µL", "+2    125 CELL/µL"
        bld_value: Masalan "0 CELL/µL", "+/-   10 CELL/µL", "+      25 CELL/µL"
    
    Returns:
        dict: {
            "WBC": {"value": "4-5", "is_red": False},
            "RBC-D": {"value": "1-2", "is_red": False},
            "RBC-U": {"value": "0-1", "is_red": False},
            "SEC": {"value": "2-3", "is_red": False},
            "TEC": {"value": "-", "is_red": False},
            "RTEC": {"value": "-", "is_red": False}
        }
    """
    
    # === LEU (Leykositlar) mapping ===
    leu_clean = leu_value.strip().upper()
    
    # WBC qiymati va qizil rang
    wbc_value = "4-5"  # Default
    wbc_red = False
    
    if "0" in leu_clean or "+-" in leu_clean or "+/-" in leu_clean:
        # 0 CELL/µL yoki +/- (15 CELL/µL)
        if "0 " in leu_clean or "0CEL" in leu_clean.replace(" ", ""):
            wbc_value = "4-5"
            wbc_red = False
        elif "15" in leu_clean:
            wbc_value = "10-15"
            wbc_red = False
    elif "+1" in leu_clean or "70" in leu_clean:
        # +1     70 CELL/µL
        wbc_value = "50-60"
        wbc_red = False
    elif "+2" in leu_clean or "125" in leu_clean:
        # +2    125 CELL/µL
        wbc_value = "100-125"
        wbc_red = True  # QIZIL!
    elif "+3" in leu_clean or "500" in leu_clean:
        # +3    500 CELL/µL
        wbc_value = "300-400"
        wbc_red = True  # QIZIL!
    
    # === BLD (Qon/Eritrositlar) mapping ===
    bld_clean = bld_value.strip().upper()
    
    # RBC-D va RBC-U qiymatlari (o'zgaruvchilar underscore bilan, lekin return qilishda tire bilan)
    rbc_d_value = "1-2"  # Default (o'zgargan)
    rbc_u_value = "0-1"  # Default (o'zgarmagan)
    rbc_red = False
    
    if "0" in bld_clean and ("CELL" in bld_clean or " 0 " in bld_clean):
        # 0 CELL/µL
        rbc_d_value = "1-2"
        rbc_u_value = "0-1"
        rbc_red = False
    elif "10" in bld_clean or "+/-" in bld_clean or "+-" in bld_clean:
        # +/-   10 CELL/µL
        rbc_d_value = "8-10"
        rbc_u_value = "1-2"
        rbc_red = False
    elif "25" in bld_clean or ("+1" in bld_clean or "+ " in bld_clean):
        # +      25 CELL/µL
        rbc_d_value = "15-20"
        rbc_u_value = "3-4"
        rbc_red = True  # QIZIL!
    elif "80" in bld_clean or "+2" in bld_clean:
        # +2      80 CELL/µL
        rbc_d_value = "40-50"
        rbc_u_value = "5-10"
        rbc_red = True  # QIZIL!
    elif "200" in bld_clean or "+3" in bld_clean:
        # +3      200 CELL/µL
        rbc_d_value = "100-120"
        rbc_u_value = "10-20"
        rbc_red = True  # QIZIL!
    
    # === Epiteliy (har doim standart) ===
    sec_value = "2-3"
    tec_value = "-"
    rtec_value = "-"
    
    return {
        "WBC": {"value": wbc_value, "is_red": wbc_red},
        "RBC-D": {"value": rbc_d_value, "is_red": rbc_red},  # Blankada "RBC-D" shaklida
        "RBC-U": {"value": rbc_u_value, "is_red": rbc_red},  # Blankada "RBC-U" shaklida
        "SEC": {"value": sec_value, "is_red": False},
        "TEC": {"value": tec_value, "is_red": False},
        "RTEC": {"value": rtec_value, "is_red": False}
    }


def parse_urit_block(block: str) -> dict:
    """
    URIT-50 dan kelgan bitta blokni lug'atga aylantiradi.
    Bundan tashqari qaysi ko'rsatkich oldida * bo'lganini ham saqlaydi.
    """
    lines = [ln.strip("\r") for ln in block.splitlines()]
    result = {}
    abnormal = set()  # * bo'lgan kodlar

    sample_no = None
    sample_id_from_urit = None  # ✅ URIT-50 dan kelgan ID (sample_id)
    sample_date = None
    sample_time = None

    # ID, NO va sana/vaqtni ajratib olish
    for line in lines:
        line_clean = line.strip()

        # ✅ ID: qismini o'qish (sample_id sifatida ishlatiladi)
        if line_clean.startswith("ID:"):
            sample_id_from_urit = line_clean.replace("ID:", "").strip()
            print(f"🆔 URIT-50 dan ID (sample_id) topildi: {sample_id_from_urit}")

        if line_clean.startswith("NO."):
            parts = line_clean.split()
            if len(parts) >= 2:
                sample_no = parts[0].replace("NO.", "")
                sample_date = parts[1]
                print(f"📅 URIT-50 dan DATE topildi: {sample_date}")

        elif ":" in line_clean and sample_time is None:
            maybe = line_clean.replace(" ", "")
            # Vaqt formatlarini tekshirish: HH:MM:SS yoki HH:MM
            if len(maybe) >= 5 and maybe[2] == ":":
                # HH:MM yoki HH:MM:SS
                if len(maybe) == 8 and maybe[5] == ":":  # HH:MM:SS
                    sample_time = maybe
                    print(f"🕐 URIT-50 dan TIME topildi: {sample_time}")
                elif len(maybe) == 5:  # HH:MM
                    sample_time = maybe
                    print(f"🕐 URIT-50 dan TIME topildi: {sample_time}")

    # Har bir analiz ko'rsatkichini ajratib olish
    for i, line in enumerate(lines):
        original = line
        starred = "*" in original            # shu qatorda * bormi?
        line_nostar = original.replace("*", "")

        for analyte in ANALYTES:
            if re.search(r"\b" + re.escape(analyte) + r"\b", line_nostar):
                parts = re.split(r"\b" + re.escape(analyte) + r"\b",
                                 line_nostar,
                                 maxsplit=1)
                raw_after = parts[1]
                after = raw_after.strip(" :-\t")

                j = i + 1
                while j < len(lines) and not any(
                    re.search(r"\b" + re.escape(a) + r"\b", lines[j])
                    for a in ANALYTES
                ):
                    after += " " + lines[j].replace("*", "").strip()
                    j += 1

                after = after.replace("CELL/uL", "CELL/µL").strip()
                # NIT (va boshqa) uchun manfiy natija: '-' belgisini saqlash
                if not after and "-" in raw_after:
                    after = "-"
                result[analyte] = after

                if starred:
                    abnormal.add(analyte)
                break

    result_all = {
        "NO": sample_no,
        "ID": sample_id_from_urit,  # ✅ URIT-50 dan kelgan ID (sample_id)
        "DATE": sample_date,
        "TIME": sample_time,
    }
    result_all.update(result)
    result_all["ABNORMAL"] = list(abnormal)
    return result_all

def optimize_page_spacing(doc):
    """
    Blankaning tepasidagi bo'sh joylarni qisqartirish va 1 sahifaga sig'adigan qilish.
    - Tepadan 0.6 sm qoldirish (foydalanuvchi talabi)
    - Jadval qatorlarini 0.4 sm balandlik qilish (foydalanuvchi talabi)
    - Paragraph spacing ni minimal qilish
    """
    from docx.shared import Pt, Cm, Inches
    
    # Sahifa margins ni to'g'rilash
    # Tepadan minimal margin - foydalanuvchi talabi: ko'p joy qolmasligi kerak
    # 0.4-0.5 sm tepadan (11-14 pt)
    for section in doc.sections:
        section.top_margin = Cm(0.3)     # Tepadan 0.3 sm (minimal - foydalanuvchi talabi: ko'p joy qolmasligi)
        section.bottom_margin = Pt(5)    # Pastdan minimal (5 pt)
        section.left_margin = Pt(28)     # Chapdan 1 sm (28 pt)
        section.right_margin = Pt(28)    # O'ngdan 1 sm (28 pt)
        
        # Header va Footer margins ni ham qisqartirish (agar bor bo'lsa)
        try:
            section.header_distance = Pt(0)  # Header dan minimal (0 pt)
            section.footer_distance = Pt(0)  # Footer dan minimal (0 pt)
        except:
            pass
    
    # Tepadagi paragraphlarning bo'sh joylarini qisqartirish
    # LЕKIN sarlavhalarni TO'LIQ saqlab qolish kerak
    for para in doc.paragraphs:
        para_text = para.text or ""
        pf = para.paragraph_format
        
        # Sarlavhalarni TO'LIQ saqlab qolish - hech qanday o'zgartirish kiritmaslik
        if any(keyword in para_text for keyword in ["FIZIK", "XOSSALARI", "MIKROSKOPIYASI", "CHO'KMASI", "SIYDIKNING", "TAXLILI"]):
            # Sarlavhalar uchun minimal spacing - faqat agar yo'q bo'lsa
            if pf.space_before is None or pf.space_before < Pt(3):
                pf.space_before = Pt(3)
            if pf.space_after is None or pf.space_after < Pt(2):
                pf.space_after = Pt(2)
            if pf.line_spacing is None:
                pf.line_spacing = 1.0
            # Sarlavha ko'rinishini TO'LIQ saqlab qolish - hech qanday o'zgartirish kiritmaslik
            continue
        
        # Agar "Test vaqti" bo'lsa, uning spacing ni set_test_time_font_13 to'g'rilaydi
        if "Test vaqti" not in para_text:
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)  # Minimal spacing (0 pt)
            pf.line_spacing = 1.0
        
        # Birinchi paragraphlarning tepasidagi joyni qisqartirish
        # Avvalgi paragraph bo'lmasa (birinchi paragraph)
        if para == doc.paragraphs[0] if doc.paragraphs else None:
            pf.space_before = Pt(0)  # Birinchi paragraphdan tepaga joy yo'q
    
    # Jadval qatorlarini 0.4 sm balandlik qilish (foydalanuvchi talabi)
    # 0.4 sm ≈ 11.34 pt, lekin Word da 1 cm = 567 twips (twentieth of a point)
    # 0.4 cm = 0.4 * 567 = 226.8 twips ≈ 227 twips
    row_height_twips = 227  # 0.4 sm
    
    for table in doc.tables:
        # Jadval ichidagi kataklarning spacing ni qisqartirish
        for row in table.rows:
            # Qator balandligini 0.4 sm qilish
            try:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls, qn
                
                # Qator balandligini o'rnatish
                tr = row._element.tr
                trPr = tr.get_or_add_trPr()
                
                # Eski trHeight ni o'chirish
                old_height = trPr.find(qn('w:trHeight'))
                if old_height is not None:
                    trPr.remove(old_height)
                
                # Yangi trHeight yaratish
                rowHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{row_height_twips}" w:hRule="exact"/>')
                trPr.append(rowHeight)
            except Exception as e:
                # Agar xato bo'lsa, faqat spacing ni qisqartirish
                print(f"   ⚠️ Jadval qator balandligini o'rnatishda xato: {e}")
            
            # Jadval ichidagi kataklarni formatlash
            for cell in row.cells:
                # Avval katak ichidagi matnni tekshirish - agar sarlavha bo'lsa, saqlab qolamiz
                cell_text = ""
                for para in cell.paragraphs:
                    cell_text += para.text + " "
                
                cell_text_upper = cell_text.upper()
                is_header = any(keyword in cell_text_upper for keyword in [
                    "FIZIK", "XOSSALARI", "MIKROSKOPIYASI", "CHO'KMASI", 
                    "SIYDIKNING", "TAXLILI", "KO'RSATKICH", "KOD", "NATIJA", "NORMA"
                ])
                
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    # Agar sarlavha bo'lsa, minimal spacing saqlaymiz
                    if is_header:
                        pf.space_before = Pt(1)
                        pf.space_after = Pt(1)
                    else:
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                    pf.line_spacing = 1.0

def normalize_doc_format(doc):
    """
    Hamma paragraph va jadval ichidagi matnlarni:
    - Times New Roman
    - 12 pt
    - interval 1.0
    - before/after = 0 pt
    qiladi.
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn

    # Oraliqlarni minimal qilish
    for para in doc.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        for run in para.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
            run.font.size = Pt(12)
            
    # Jadval ichidagi hamma kataklar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    pf = para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0

                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                        run.font.size = Pt(12)




def set_test_time_font_13(doc):
    """'Test vaqti:' qatori 13 pt bo'lishi uchun va jadvaldan minimal pastda."""
    from docx.shared import Pt
    
    target = "Test vaqti"
    
    # Jadvaldan keyingi "Test vaqti" paragraphni topish
    # Python-docx da jadval va paragraphlar ketma-ket bo'ladi
    # Demak, jadvaldan keyingi paragraphni topish uchun barcha elementlarni tekshirish kerak
    
    # Avval jadval ichida bo'lsa
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if target in (para.text or ""):
                        # Jadval ichida bo'lsa, fontni o'rnatamiz va QALIN qilamiz
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(13)
                            run.bold = True  # ✅ Test vaqti qalin bo'lishi kerak
    
    # Oddiy paragraphlarda "Test vaqti" ni qidirish
    # Jadvaldan keyingi paragraph bo'lishi ehtimoli yuqori
    for para_idx, para in enumerate(doc.paragraphs):
        if target in (para.text or ""):
            # Font 13 pt va QALIN (bold)
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.bold = True  # ✅ Test vaqti qalin bo'lishi kerak
            
            # Jadvaldan minimal pastda bo'lishi uchun (foydalanuvchi talabi: ko'p joy qolmasligi)
            pf = para.paragraph_format
            pf.space_before = Pt(3)  # Jadvaldan pastda minimal (3 pt = 0.1 sm)
            pf.space_after = Pt(0)

from docx.oxml.ns import qn

def format_results_only(doc):
    """
    FIZIK–KIMYOVIY XOSSALARI dan PASTINI formatlaydi.
    Sarlavhalarni TO'LIQ saqlab qoladi - hech qanday o'zgartirish kiritmaydi.
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    start = False

    # Oddiy paragraphlar
    for para in doc.paragraphs:
        para_text = para.text or ""
        
        # Sarlavha bo'lsa, uni TO'LIQ saqlab qolamiz - hech qanday o'zgartirish kiritmaslik
        if "FIZIK" in para_text and "XOSSALARI" in para_text:
            # Sarlavhani to'liq saqlab qolish - faqat spacing ni minimal qilamiz
            pf = para.paragraph_format
            # Spacing ni minimal qilamiz, lekin ko'rinishini saqlaymiz
            if pf.space_before is None or pf.space_before < Pt(4):
                pf.space_before = Pt(4)  # Minimal joy
            if pf.space_after is None or pf.space_after < Pt(2):
                pf.space_after = Pt(2)   # Minimal joy
            
            # Sarlavha matnini TO'LIQ saqlab qolish - hech qanday o'zgartirish kiritmaslik
            # Faqat agar format yo'q bo'lsa, qo'shamiz
            for run in para.runs:
                # Agar run matni bo'sh bo'lsa, o'tkazib yuboramiz
                if not run.text.strip():
                    continue
                # Formatni saqlab qolish - faqat zarur bo'lsa qo'shamiz
                if not run.font.name:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                if not run.font.size:
                    run.font.size = Pt(13)  # Sarlavha biroz kattaroq
                if run.bold is None:
                    run.bold = True  # Sarlavha qalin bo'lishi kerak
            
            start = True
            continue  # Sarlavhani formatlashdan keyin keyingi paragraphga o'tamiz

        if not start:
            continue

        # Sarlavha bo'lmagan paragraphlar uchun
        # Qatorlar orasi qisqa qilish
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0

        for run in para.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
            run.font.size = Pt(12)

    # Jadval ichidagi natijalar
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Katak ichidagi matnni tekshirish - agar sarlavha bo'lsa, saqlab qolamiz
                cell_text = ""
                for para in cell.paragraphs:
                    cell_text += para.text + " "
                
                cell_text_upper = cell_text.upper()
                is_header_row = any(keyword in cell_text_upper for keyword in [
                    "FIZIK", "XOSSALARI", "MIKROSKOPIYASI", "CHO'KMASI",
                    "KO'RSATKICH", "KOD", "NATIJA", "NORMA", "O'LCHOV BIRLIGI"
                ])
                
                for para in cell.paragraphs:
                    if start:
                        # Qatorlar orasi qisqa qilish
                        pf = para.paragraph_format
                        # Agar sarlavha qatori bo'lsa, minimal spacing saqlaymiz
                        if is_header_row:
                            pf.space_before = Pt(2)
                            pf.space_after = Pt(2)
                        else:
                            pf.space_before = Pt(0)
                            pf.space_after = Pt(0)
                        pf.line_spacing = 1.0
                        
                        for run in para.runs:
                            run.font.name = "Times New Roman"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                            run.font.size = Pt(12)
                            # Sarlavha qatorlari qalin bo'lishi kerak
                            if is_header_row:
                                run.bold = True


# ====== DATABASE FUNCTIONS ======
# ✅ Connection Pooling - WiFi da tezroq ishlash uchun
_db_connection = None
_db_connection_lock = threading.Lock()

def db_conn():
    """
    MySQL ga ulanish - Connection Pooling bilan
    WiFi da sekin ishlash muammosini hal qiladi
    """
    global _db_connection
    
    with _db_connection_lock:
        # Agar ulanish mavjud va ishlayotgan bo'lsa, qaytaramiz
        if _db_connection is not None:
            try:
                # Ulanish ishlayotganini tekshirish
                _db_connection.ping(reconnect=True, attempts=1, delay=0)
                return _db_connection
            except:
                # Agar ulanish yopilgan bo'lsa, None qilamiz
                _db_connection = None
        
        # Yangi ulanish yaratish
        try:
            # Connection timeout va pool sozlamalari
            config = DB_CONFIG.copy()
            config.update({
                "connection_timeout": 5,  # 5 soniya timeout
                "autocommit": False,  # Transaction uchun
                "pool_reset_session": True,  # Pool reset
            })
            
            _db_connection = mysql.connector.connect(**config)
            print("✅ Database ulanishi yaratildi (Connection Pooling)")
            return _db_connection
        except Exception as e:
            print(f"❌ Ma'lumotlar bazasiga ulana olmadim: {e}")
            print(f"   IP: {DB_CONFIG['host']}")
            print(f"   Database: {DB_CONFIG['database']}")
            _db_connection = None
            return None

def db_close():
    """Database ulanishini yopish (dastur yakunlanganda)"""
    global _db_connection
    with _db_connection_lock:
        if _db_connection is not None:
            try:
                _db_connection.close()
                print("✅ Database ulanishi yopildi")
            except:
                pass
            finally:
                _db_connection = None

def get_patient_by_barcode(barcode: str):
    """
    Barcode orqali bemor ma'lumotlarini olish
    Barcode: kod_yollanma, natija_kodi yoki sample_id
    Qidirish variantlari:
    1. To'liq mos kelishi (sample_id = barcode) - ENG YUQORI USTUNLIK
    2. To'liq mos kelishi (kod_yollanma = barcode)
    3. To'liq mos kelishi (natija_kodi = barcode)
    4. Sample_id ichida barcode mavjud (LIKE)
    5. Kod_yollanma ichida barcode mavjud (LIKE)
    6. Natija_kodi ichida barcode mavjud (LIKE)
    7. Barcode oxirgi qismi (oxirgi 10-15 raqam)
    """
    conn = db_conn()
    if not conn:
        return None
    
    try:
        cursor = conn.cursor(dictionary=True)
        
        # Barcode ni tozalash
        barcode_clean = barcode.strip()
        print(f"🔍 Barcode qidiryapman: '{barcode_clean}'")
        
        # bemorlar jadvalidan qidirish - sample_id bilan ham qidirish
        query = """
            SELECT 
                id,
                sample_id,
                kod_yollanma,
                fish,
                yosh,
                jins,
                tugilgan_sana,
                telefon,
                shifokor,
                natija_kodi
            FROM bemorlar
            WHERE sample_id = %s
               OR kod_yollanma = %s 
               OR natija_kodi = %s
               OR sample_id LIKE %s
               OR kod_yollanma LIKE %s
               OR natija_kodi LIKE %s
               OR sample_id LIKE %s
               OR kod_yollanma LIKE %s
               OR natija_kodi LIKE %s
               OR sample_id LIKE %s
               OR kod_yollanma LIKE %s
               OR natija_kodi LIKE %s
            ORDER BY 
                CASE 
                    WHEN sample_id = %s THEN 1
                    WHEN kod_yollanma = %s THEN 2
                    WHEN natija_kodi = %s THEN 3
                    WHEN sample_id LIKE %s THEN 4
                    WHEN kod_yollanma LIKE %s THEN 5
                    WHEN natija_kodi LIKE %s THEN 6
                    ELSE 7
                END,
                id DESC
            LIMIT 1
        """
        
        # Qidirish variantlari
        barcode_full = barcode_clean
        barcode_contains = f"%{barcode_clean}%"
        barcode_end_15 = f"%{barcode_clean[-15:]}" if len(barcode_clean) >= 15 else f"%{barcode_clean}"
        barcode_end_10 = f"%{barcode_clean[-10:]}" if len(barcode_clean) >= 10 else f"%{barcode_clean}"
        
        cursor.execute(query, (
            barcode_full,  # sample_id = barcode (ENG YUQORI USTUNLIK)
            barcode_full,  # kod_yollanma = barcode
            barcode_full,  # natija_kodi = barcode
            barcode_contains,  # sample_id LIKE %barcode%
            barcode_contains,  # kod_yollanma LIKE %barcode%
            barcode_contains,  # natija_kodi LIKE %barcode%
            barcode_end_15,  # sample_id LIKE %oxirgi15%
            barcode_end_15,  # kod_yollanma LIKE %oxirgi15%
            barcode_end_15,  # natija_kodi LIKE %oxirgi15%
            barcode_end_10,  # sample_id LIKE %oxirgi10%
            barcode_end_10,  # kod_yollanma LIKE %oxirgi10%
            barcode_end_10,  # natija_kodi LIKE %oxirgi10%
            # ORDER BY uchun (ustunlik tartibi)
            barcode_full,  # priority 1 - sample_id to'liq mos
            barcode_full,  # priority 2 - kod_yollanma to'liq mos
            barcode_full,  # priority 3 - natija_kodi to'liq mos
            barcode_contains,  # priority 4 - sample_id LIKE
            barcode_contains,  # priority 5 - kod_yollanma LIKE
            barcode_contains   # priority 6 - natija_kodi LIKE
        ))
        result = cursor.fetchone()
        
        if result:
            print(f"✅ Bemor topildi: {result['fish']}")
            print(f"   ID: {result['id']}")
            print(f"   Sample ID: {result.get('sample_id', 'NULL')}")
            print(f"   Kod yo'llanma: {result['kod_yollanma']}")
            print(f"   Natija kodi: {result['natija_kodi']}")
        else:
            print(f"⚠️ Bemor topilmadi: {barcode}")
            print(f"   Tekshirilgan variantlar:")
            print(f"   - To'liq sample_id = '{barcode_clean}'")
            print(f"   - To'liq kod_yollanma = '{barcode_clean}'")
            print(f"   - To'liq natija_kodi = '{barcode_clean}'")
            print(f"   - sample_id LIKE '%{barcode_clean}%'")
            print(f"   - kod_yollanma LIKE '%{barcode_clean}%'")
            print(f"   - natija_kodi LIKE '%{barcode_clean}%'")
        
        return result
        
    except Exception as e:
        print(f"❌ Baza xatosi: {e}")
        import traceback
        traceback.print_exc()
        # Xato bo'lsa, ulanishni qayta tiklashga harakat qilamiz
        global _db_connection
        with _db_connection_lock:
            if _db_connection:
                try:
                    _db_connection.close()
                except:
                    pass
                _db_connection = None
        return None
    # ✅ Connection ni yopmaymiz - pool da saqlaymiz

def save_result_to_db(sample_no: str, parsed_data: dict, patient_data: dict = None, word_path: str = ""):
    """
    Natijalarni bazaga saqlash
    """
    conn = db_conn()
    if not conn:
        return None
    
    try:
        cursor = conn.cursor()
        
        # Bemor ID ni topish
        bemor_id = None
        if patient_data:
            bemor_id = patient_data.get('id')
        elif sample_no:
            # Sample NO orqali bemor topish - sample_id ga ustunlik beramiz
            # Avval sample_id orqali qidiramiz (to'liq mos kelishi)
            cursor.execute("""
                SELECT id FROM bemorlar 
                WHERE sample_id = %s 
                   OR sample_id LIKE %s
                   OR kod_yollanma LIKE %s 
                   OR natija_kodi LIKE %s
                ORDER BY 
                    CASE 
                        WHEN sample_id = %s THEN 1
                        WHEN sample_id LIKE %s THEN 2
                        ELSE 3
                    END
                LIMIT 1
            """, (
                sample_no,  # sample_id = sample_no (to'liq mos)
                f"%{sample_no}%",  # sample_id LIKE %sample_no%
                f"%{sample_no}%",  # kod_yollanma LIKE %sample_no%
                f"%{sample_no}%",  # natija_kodi LIKE %sample_no%
                # ORDER BY uchun
                sample_no,  # priority 1
                f"%{sample_no}%"  # priority 2
            ))
            row = cursor.fetchone()
            if row:
                bemor_id = row[0]
                print(f"   ✅ Bemor topildi sample_id orqali: ID={bemor_id}, sample_no='{sample_no}'")
        
        if not bemor_id:
            print("⚠️ Bemor ID topilmadi, natija saqlanmadi")
            return None
        
        # Buyurtma topish
        cursor.execute("SELECT id FROM orders WHERE bemor_id = %s ORDER BY sana_vaqt DESC LIMIT 1", (bemor_id,))
        order_row = cursor.fetchone()
        order_id = order_row[0] if order_row else None
        
        if not order_id:
            print("⚠️ Buyurtma topilmadi, natija saqlanmadi")
            return None
        
        # ✅ Natijalarni 'results' jadvaliga saqlash
        print(f"💾 Natijalarni bazaga saqlayapman...")
        print(f"   Bemor ID: {bemor_id}")
        print(f"   Order ID: {order_id}")
        
        # Avval results jadvalidan topish yoki yaratish
        cursor.execute("""
            SELECT id FROM results 
            WHERE order_id = %s AND test_turi = %s
            LIMIT 1
        """, (order_id, 'SIYDIK_TAHLILI'))
        
        existing_result = cursor.fetchone()
        
        if existing_result:
            # Agar mavjud bo'lsa, yangilaymiz
            result_id = existing_result[0]
            cursor.execute("""
                UPDATE results 
                SET natija_fayl = %s, sana_vaqt = NOW(), updated_at = NOW()
                WHERE id = %s
            """, (word_path, result_id))
            print(f"   Mavjud natija yangilandi: Result ID={result_id}")
        else:
            # Yangi natija yaratamiz
            cursor.execute("""
                INSERT INTO results (order_id, bemor_id, test_turi, natija_fayl, sana_vaqt, created_at)
                VALUES (%s, %s, %s, %s, NOW(), NOW())
            """, (order_id, bemor_id, 'SIYDIK_TAHLILI', word_path))
            result_id = cursor.lastrowid
            print(f"   Yangi natija yaratildi: Result ID={result_id}")
        
        # Result_items jadvalidagi eski natijalarni o'chirish (agar mavjud bo'lsa)
        cursor.execute("DELETE FROM result_items WHERE result_id = %s", (result_id,))
        
        # Result_items jadvaliga har bir analiz natijasini saqlash
        items_count = 0
        for code, value in parsed_data.items():
            if code in ['NO', 'DATE', 'TIME', 'ABNORMAL']:
                continue
            
            if not value:  # Bo'sh qiymatni o'tkazib yuboramiz
                continue
            
            # Analiz kodini saqlash
            cursor.execute("""
                INSERT INTO result_items (result_id, test_kodi, natija, created_at)
                VALUES (%s, %s, %s, NOW())
            """, (result_id, code, str(value)))
            items_count += 1
        
        # ✅ pdf_documents jadvaliga ham saqlash (monoblok dastur uchun)
        if word_path:
            try:
                # Avval mavjud yozuvni tekshirish
                cursor.execute("""
                    SELECT id FROM pdf_documents 
                    WHERE order_id = %s AND file_type = 'SIYDIK_TAHLILI'
                    LIMIT 1
                """, (order_id,))
                existing_pdf = cursor.fetchone()
                
                if existing_pdf:
                    # Yangilash
                    cursor.execute("""
                        UPDATE pdf_documents 
                        SET file_path = %s, created_at = NOW()
                        WHERE id = %s
                    """, (word_path, existing_pdf[0]))
                    print(f"   ✅ pdf_documents yangilandi: ID={existing_pdf[0]}")
                else:
                    # Yangi yozuv yaratish
                    # sample_id ni olish
                    sample_id = None
                    if patient_data:
                        sample_id = patient_data.get('sample_id', '')
                    if not sample_id:
                        # orders jadvalidan sample_id ni olish
                        cursor.execute("SELECT sample_id FROM orders WHERE id = %s", (order_id,))
                        order_row = cursor.fetchone()
                        if order_row:
                            sample_id = order_row[0]
                    
                    cursor.execute("""
                        INSERT INTO pdf_documents (order_id, sample_id, file_path, file_type, created_at)
                        VALUES (%s, %s, %s, 'SIYDIK_TAHLILI', NOW())
                    """, (order_id, sample_id, word_path))
                    pdf_id = cursor.lastrowid
                    print(f"   ✅ pdf_documents yaratildi: ID={pdf_id}")
            except Exception as e:
                print(f"   ⚠️ pdf_documents ga saqlashda xato: {e}")
                # Xato bo'lsa ham davom etamiz
        
        print(f"✅ Natijalar bazaga saqlandi!")
        print(f"   Result ID: {result_id}")
        print(f"   Natija elementlari: {items_count} ta")
        print(f"   Fayl: {word_path}")
        
        conn.commit()
        return order_id
        
    except Exception as e:
        print(f"❌ Bazaga saqlashda xatolik: {e}")
        if conn:
            conn.rollback()
        # Xato bo'lsa, ulanishni qayta tiklashga harakat qilamiz
        global _db_connection
        with _db_connection_lock:
            if _db_connection:
                try:
                    _db_connection.close()
                except:
                    pass
                _db_connection = None
        return None
    # ✅ Connection ni yopmaymiz - pool da saqlaymiz

# ====== SCANNER INPUT ======
def barcode_scanner_thread():
    """
    Skaner orqali barcode o'qish (threading orqali)
    Skaner odatda keyboard input sifatida ishlaydi
    """
    global current_barcode
    
    print("Skaner tayyor. Barcode skanerlang...")
    
    while True:
        try:
            # Skaner matnni yuboradi va Enter bosadi
            barcode = input().strip()
            
            if barcode:
                current_barcode = barcode
                barcode_queue.put(barcode)
                print(f"📷 Barcode qabul qilindi: {barcode}")
        except EOFError:
            break
        except Exception as e:
            print(f"❌ Skaner xatosi: {e}")

# ====== MIKROSKOPIYA NATIJALARINI AVTOMATIK TO'LDIRISH ======
def fill_microscopy_results(doc: Document, data: dict, abnormal_codes: set):
    """
    2-rasmdagi jadvalga qarab mikroskopiya qismini avtomatik to'ldiradi.
    LEU, BLD, NIT natijalariga qarab mikroskopiya qismini to'ldiradi.
    """
    from docx.oxml.ns import qn
    
    def parse_value(value_str: str):
        """URIT-50 dan kelgan qiymatni parse qilish"""
        if not value_str:
            return None, 0
        
        value_str = value_str.strip().upper()
        
        # *+1, *+2, *+3, +/-, *+- formatlarni aniqlash
        has_star = "*" in value_str
        has_plus_minus = "+/-" in value_str or "+-" in value_str or "*+-" in value_str
        has_plus = "+" in value_str and not has_plus_minus
        
        # Raqamni ajratish
        import re
        numbers = re.findall(r'\d+', value_str)
        numeric_value = int(numbers[0]) if numbers else 0
        
        # Format aniqlash
        if has_plus_minus:
            return " +/-", numeric_value
        elif "+1" in value_str or (has_plus and numeric_value > 0 and numeric_value < 50):
            return "+1", numeric_value
        elif "+2" in value_str or (has_plus and numeric_value >= 50 and numeric_value < 150):
            return "+2", numeric_value
        elif "+3" in value_str or (has_plus and numeric_value >= 150):
            return "+3", numeric_value
        elif numeric_value == 0:
            return "0", 0
        else:
            return "", numeric_value
    
    def find_row_by_text(table, search_texts: list, col_index: int = 0):
        """Jadvalda matn bo'yicha qatorni topish (barcha ustunlarni tekshiradi)"""
        for row in table.rows:
            # Barcha kataklarni tekshirish (mikroskopiya qismida kod boshqa ustunda bo'lishi mumkin)
            for col_idx in range(min(len(row.cells), 5)):  # Faqat birinchi 5 ustunni tekshiramiz
                if len(row.cells) > col_idx:
                    cell_text = row.cells[col_idx].text.strip().upper()
                    for search in search_texts:
                        if search.upper() in cell_text:
                            return row
        return None
    
    def fill_cell_with_text(cell, text: str, font_size: int = 12, is_red: bool = False):
        """Katakni matn bilan to'ldirish - formatni buzmasdan (qizil yoki qora)"""
        # ✅ Formatni buzmasdan tozalash - faqat matnni o'chiramiz
        for para in cell.paragraphs:
            # Barcha runlarni tozalaymiz, lekin paragraph strukturini saqlaymiz
            for run in para.runs:
                run.text = ""
            
            # Agar paragraph bo'sh bo'lsa, birinchi paragraphni tozalaymiz
            if len(para.runs) == 0:
                para.add_run("")
        
        # Birinchi paragraphni olamiz
        para = cell.paragraphs[0]
        para_format = para.paragraph_format
        para_format.space_before = Pt(0)
        para_format.space_after = Pt(0)
        para_format.line_spacing = 1.0
        
        # Birinchi runni topamiz yoki yangi run qo'shamiz
        if len(para.runs) > 0:
            run = para.runs[0]
            run.text = text
        else:
            run = para.add_run(text)
        
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
        run.font.size = Pt(font_size)
        run.bold = True
        # Qizil yoki qora rang
        if is_red:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Qizil rang
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Qora rang
    
    def fill_cell_with_red_text(cell, text: str, font_size: int = 12):
        """Katakni qizil matn bilan to'ldirish (eski funksiya - orqaga moslik uchun)"""
        fill_cell_with_text(cell, text, font_size, is_red=True)
    
    print("\n🔬 Mikroskopiya qismini to'ldiryapman...")
    
    # 1. LEU (Leykotsitlar) -> WBC (Leykotsitlar)
    leu_value = data.get("LEU", "")
    if leu_value:
        format_type, numeric_val = parse_value(leu_value)
        print(f"   LEU: '{leu_value}' -> format: '{format_type}', value: {numeric_val}")
        
        # 2-rasmdagi jadvalga qarab WBC qiymatini aniqlash
        wbc_result = None
        if format_type == "0" or numeric_val == 0:
            wbc_result = "4-5"
        elif format_type == " +/-" or (format_type == "" and 10 <= numeric_val <= 20):
            wbc_result = "10-15"  # +/- 15 CELL/uL -> 10-15
        elif format_type == "+1" or (format_type == "" and 50 <= numeric_val < 100):
            wbc_result = "50-60"  # +1 70 CELL/uL -> 50-60
        elif format_type == "+2" or (format_type == "" and 100 <= numeric_val < 200):
            wbc_result = "100-125"  # +2 125 CELL/uL -> 100-125
        elif format_type == "+3" or (format_type == "" and numeric_val >= 200):
            wbc_result = "300-400"  # +3 500 CELL/uL -> 300-400
        
        if wbc_result:
            # WBC qatorini topish (mikroskopiya qismida)
            found_wbc = False
            for table in doc.tables:
                # Avval "MIKROSKOPIYASI" yoki "CHO'KMASI" so'zini qidiramiz
                is_microscopy_table = False
                for row in table.rows:
                    for cell in row.cells:
                        cell_text_upper = cell.text.upper()
                        if any(word in cell_text_upper for word in ["MIKROSKOPIYASI", "CHO'KMASI", "MİKROSKOPİYASI", "CHO'KMAS"]):
                            is_microscopy_table = True
                            break
                    if is_microscopy_table:
                        break
                
                if not is_microscopy_table:
                    continue
                
                print(f"   📋 Mikroskopiya jadvali topildi, WBC qatorini qidiryapman...")
                
                # ✅ WBC qatorini topish - faqat mikroskopiya jadvalidagi Leykotsitlar qatorini topish
                # Avval Ko'rsatkich ustunida "Leykotsitlar" borligini tekshirish
                wbc_row = None
                for row in table.rows:
                    # Sarlavha qatorlarini o'tkazib yuborish
                    row_text = ""
                    for cell in row.cells:
                        row_text += cell.text + " "
                    row_text_upper = row_text.upper()
                    if any(keyword in row_text_upper for keyword in ["KO'RSATKICH", "KOD", "NATIJA", "NORMA", "MIKROSKOPIYASI", "CHO'KMASI"]):
                        continue
                    
                    # Ko'rsatkich ustunida "Leykotsitlar" va kod ustunida "WBC" bo'lishi kerak
                    if len(row.cells) > 0:
                        indicator = row.cells[0].text.strip().upper()
                        if "LEYKOTSITLAR" in indicator or "LEYUKOTSITLAR" in indicator:
                            # Kod ustunini tekshirish
                            if len(row.cells) > 1:
                                code_cell = row.cells[1].text.strip().upper()
                                if "WBC" in code_cell or (len(row.cells) > 2 and "WBC" in row.cells[2].text.strip().upper()):
                                    wbc_row = row
                                    print(f"   ✅ WBC qatori topildi (mikroskopiya): Ko'rsatkich='{row.cells[0].text}', Kod='{code_cell}'")
                                    break
                
                # Agar topilmasa, eski usul bilan qidirish
                if not wbc_row:
                    wbc_row = find_row_by_text(table, [
                        "ЛЕЙКОЦИТЛАР", "LEYKOTSITLAR", "WBC", "Лейк", "Leykotsitlar",
                        "ЛЕЙКОЦИТЫ", "LEUKOCYTES", "Лейкоциты"
                    ])
                
                if wbc_row and len(wbc_row.cells) >= 3:
                    print(f"   ✅ WBC qatori topildi, {len(wbc_row.cells)} ta katak")
                    # Jadval strukturasini aniqlash
                    # Odatda: 0=Ko'rsatkich, 1=Kod, 2=Natija, 3=Norma, 4=O'lchov birligi
                    
                    # Har bir katakni ko'rsatish
                    for col_idx in range(min(len(wbc_row.cells), 5)):
                        print(f"      Katak {col_idx}: '{wbc_row.cells[col_idx].text}'")
                    
                    # WBC kodini topish va natija ustunini aniqlash
                    result_cell = None
                    wbc_code_col = None
                    
                    # Kod ustunini topish (ustun 1 yoki 2)
                    # ✅ EHTIYOT: Agar "Kod" ustunida "10-15" bo'lsa, uni tozalash kerak
                    for col_idx in range(min(len(wbc_row.cells), 5)):
                        cell_text = wbc_row.cells[col_idx].text.strip()
                        cell_text_upper = cell_text.upper()
                        print(f"      Katak {col_idx} tekshirilmoqda: '{cell_text}'")
                        
                        # Agar "Kod" ustunida "10-15" yoki boshqa raqam bo'lsa, uni tozalash
                        if col_idx == 1 or col_idx == 2:  # Kod ustunlari
                            if "10-15" in cell_text or "10-" in cell_text or "-15" in cell_text:
                                print(f"      ⚠️ Kod ustunida noto'g'ri qiymat topildi: '{cell_text}', tozalayapman...")
                                # Kod ustunini tozalash - faqat "WBC" ni qoldirish
                                for para in wbc_row.cells[col_idx].paragraphs:
                                    for run in para.runs:
                                        run.text = ""
                                    if len(para.runs) == 0:
                                        para.add_run("")
                                # "WBC" ni qo'yish, agar yo'q bo'lsa
                                if "WBC" not in cell_text_upper:
                                    wbc_row.cells[col_idx].paragraphs[0].runs[0].text = "WBC"
                                else:
                                    wbc_row.cells[col_idx].paragraphs[0].runs[0].text = "WBC"
                                print(f"      ✅ Kod ustuni tozalandi: 'WBC'")
                                cell_text = "WBC"  # Tozalangan qiymat
                                cell_text_upper = "WBC"
                        
                        # WBC kodini topish (to'liq so'z yoki qismi)
                        if "WBC" == cell_text_upper or "WBC" in cell_text_upper:
                            wbc_code_col = col_idx
                            print(f"      ✅ WBC kod topildi: ustun {col_idx} ('{cell_text}')")
                            break
                    
                    # Natija ustunini aniqlash
                    # Jadval strukturasiga qarab: 0=Ko'rsatkich, 1=Kod, 2=Natija, 3=Norma, 4=O'lchov birligi
                    result_col_idx = None
                    if wbc_code_col is not None:
                        # Kod topilgan ustundan keyingi ustun - Natija ustuni
                        result_col_idx = wbc_code_col + 1
                        if result_col_idx < len(wbc_row.cells):
                            result_cell = wbc_row.cells[result_col_idx]
                            print(f"      ✅ Natija ustuni topildi: ustun {result_col_idx} (koddan keyingi)")
                        else:
                            print(f"      ⚠️ Natija ustuni topilmadi (jadval chegarasidan tashqarida: {result_col_idx} >= {len(wbc_row.cells)})")
                    
                    # Agar kod topilmasa yoki natija ustuni topilmasa, strukturaga qarab aniqlaymiz
                    if not result_cell:
                        # Jadval strukturasiga qarab: 0=Ko'rsatkich, 1=Kod, 2=Natija
                        if len(wbc_row.cells) >= 3:
                            # Odatda ustun 2 - Natija
                            result_cell = wbc_row.cells[2]
                            result_col_idx = 2
                            print(f"      ⚠️ Strukturaga qarab, 2-ustunni (Natija) to'ldiryapman")
                        elif len(wbc_row.cells) >= 4:
                            result_cell = wbc_row.cells[3]
                            result_col_idx = 3
                            print(f"      ⚠️ Strukturaga qarab, 3-ustunni (Natija) to'ldiryapman")
                    
                    if result_cell:
                        # Eski natijani ko'rsatish
                        old_value = result_cell.text.strip()
                        # Normada bo'lsa (4-5), qizil emas, qora rangda
                        # Normadan tashqarida bo'lsa (10-15, 50-60, 100-125, 300-400), qizil rangda
                        is_abnormal = wbc_result != "4-5"
                        fill_cell_with_text(result_cell, wbc_result, is_red=is_abnormal)
                        color_info = "qizil" if is_abnormal else "qora"
                        col_info = f"ustun {result_col_idx}" if result_col_idx is not None else "Natija ustuni"
                        print(f"   ✅ WBC (Leykotsitlar) natija yozildi: '{old_value}' -> '{wbc_result}' ({color_info}, {col_info})")
                        found_wbc = True
                        break
                    else:
                        print(f"   ❌ Natija katagi topilmadi!")
                else:
                    print(f"   ⚠️ WBC qatori topilmadi")
            
            if not found_wbc:
                print(f"   ❌ WBC qatori topilmadi, mikroskopiya jadvalida qidirildi")
    
    # 2. BLD (Qon) -> RBC-D (o'zgargan) va RBC-U (o'zgarmagan)
    bld_value = data.get("BLD", "")
    if bld_value:
        leu_value = data.get("LEU", "0 CELL/µL")
        # get_microscopy_values funksiyasidan foydalanish
        micro_values = get_microscopy_values(leu_value, bld_value)
        
        print(f"   BLD: '{bld_value}' -> mikroskopiya qiymatlari hisoblandi")
        
        # RBC-D va RBC-U kodlarini qidirish
        for table in doc.tables:
            # Mikroskopiya jadvalini aniqlash
            is_microscopy_table = False
            for row in table.rows:
                for cell in row.cells:
                    if any(word in cell.text.upper() for word in ["MIKROSKOPIYASI", "CHO'KMASI", "MİKROSKOPİYASI"]):
                        is_microscopy_table = True
                        break
                if is_microscopy_table:
                    break
            
            if not is_microscopy_table:
                continue
            
            print(f"   📋 Mikroskopiya jadvali topildi, RBC kodlarini qidiryapman...")
            
            # ✅ YANGI YONDASHUV: Faqat KOD ustunida RBC-D va RBC-U ni qidirish
            # Jadval strukturasi: Ko'rsatkich | Kod | Natija | Norma | O'lchov birligi
            
            for row in table.rows:
                # Sarlavha qatorlarini o'tkazib yuborish
                if len(row.cells) == 0:
                    continue
                row_text_upper = row.cells[0].text.upper()
                if any(keyword in row_text_upper for keyword in ["KO'RSATKICH", "KOD", "NATIJA", "NORMA", "MIKROSKOPIYASI", "CHO'KMASI"]):
                    continue
                
                # Jadvalda kamida 3 ta ustun bo'lishi kerak
                if len(row.cells) < 3:
                    continue
                
                # ====== RBC-D (o'zgargan) ni qidirish ======
                # 1-ustunda (Kod ustuni) "RBC-D" borligini tekshirish
                kod_cell_text = row.cells[1].text.strip().upper()
                
                # RBC-D topilsa (blankada "RBC-D" shaklida yozilgan)
                if "RBC-D" in kod_cell_text:
                    print(f"   ✅ RBC-D kodi topildi: Kod ustuni='{row.cells[1].text}'")
                    
                    # get_microscopy_values dan qiymatni olish
                    if "RBC-D" in micro_values:
                        value_info = micro_values["RBC-D"]
                        rbc_d_value = value_info["value"]
                        is_abnormal_d = value_info["is_red"]
                        
                        # Natija ustuni (2-ustun)ga yozish
                        result_cell = row.cells[2]
                        old_value = result_cell.text.strip()
                        
                        # Qiymat va rang
                        fill_cell_with_text(result_cell, rbc_d_value, is_red=is_abnormal_d)
                        color_info = "qizil" if is_abnormal_d else "qora"
                        print(f"   ✅ RBC-D (o'zgargan) natija yozildi: '{old_value}' -> '{rbc_d_value}' ({color_info}, 2-ustun)")
                
                # ====== RBC-U (o'zgarmagan) ni qidirish ======
                # 1-ustunda (Kod ustuni) "RBC-U" borligini tekshirish
                elif "RBC-U" in kod_cell_text:
                    print(f"   ✅ RBC-U kodi topildi: Kod ustuni='{row.cells[1].text}'")
                    
                    # get_microscopy_values dan qiymatni olish
                    if "RBC-U" in micro_values:
                        value_info = micro_values["RBC-U"]
                        rbc_u_value = value_info["value"]
                        is_abnormal_u = value_info["is_red"]
                        
                        # Natija ustuni (2-ustun)ga yozish
                        result_cell = row.cells[2]
                        old_value = result_cell.text.strip()
                        
                        # Qiymat va rang
                        fill_cell_with_text(result_cell, rbc_u_value, is_red=is_abnormal_u)
                        color_info = "qizil" if is_abnormal_u else "qora"
                        print(f"   ✅ RBC-U (o'zgarmagan) natija yozildi: '{old_value}' -> '{rbc_u_value}' ({color_info}, 2-ustun)")
    
    print("✅ Eritrositlar qismi to'ldirildi!")

    
    # 3. NIT (Nitrit) -> Bakteriyalar
    nit_value = data.get("NIT", "")
    if nit_value and ("+" in nit_value.upper() or "ПОЛОЖИТЕЛЬНЫЙ" in nit_value.upper() or "MANFIY EMAS" in nit_value.upper()):
        print(f"   NIT: '{nit_value}' -> Bakteriyalar: +++")
        # Bakteriyalar qatorini topish
        for table in doc.tables:
            # Mikroskopiya jadvalini aniqlash
            is_microscopy_table = False
            for row in table.rows:
                for cell in row.cells:
                    if any(word in cell.text.upper() for word in ["MIKROSKOPIYASI", "CHO'KMASI", "MİKROSKOPİYASI"]):
                        is_microscopy_table = True
                        break
                if is_microscopy_table:
                    break
            
            if not is_microscopy_table:
                continue
            
            bacteria_row = find_row_by_text(table, 
                ["БАКТЕРИЯЛАР", "BAKTERIYALAR", "BACT", "Bakteriyalar", "Бактерии", "БАКТЕРИИ"])
            if bacteria_row and len(bacteria_row.cells) >= 3:
                # BACT kodini topish va natija ustunini aniqlash
                result_cell = None
                bact_code_col = None
                
                # Kod ustunini topish (ustun 1 yoki 2)
                for col_idx in range(1, min(len(bacteria_row.cells), 4)):
                    cell_text_upper = bacteria_row.cells[col_idx].text.upper().strip()
                    if "BACT" in cell_text_upper or "BAKTERIYALAR" in cell_text_upper or "БАКТЕРИ" in cell_text_upper:
                        bact_code_col = col_idx
                        print(f"      ✅ BACT kod topildi (Bakteriyalar): ustun {col_idx}")
                        break
                
                # Natija ustunini aniqlash (kod ustunidan keyingi)
                if bact_code_col is not None:
                    result_col_idx = bact_code_col + 1
                    if result_col_idx < len(bacteria_row.cells):
                        result_cell = bacteria_row.cells[result_col_idx]
                
                # Agar kod topilmasa yoki natija ustuni topilmasa, strukturaga qarab
                if not result_cell:
                    # Odatda: 0=Ko'rsatkich, 1=Kod, 2=Natija
                    if len(bacteria_row.cells) >= 3:
                        result_cell = bacteria_row.cells[2]  # Natija ustuni
                
                if result_cell:
                    fill_cell_with_red_text(result_cell, "+++")
                    print(f"   ✅ Bakteriyalar: +++ (qizil, Natija ustuniga)")
                break
    
    print("✅ Mikroskopiya qismi to'ldirildi\n")


# ====== WORD BLANKAGA YOZISH ======
def create_doc_from_data(data: dict, fish: str = "", yosh: str = "", tugilgan_sana: str = "", patient_data: dict = None) -> str:
    """
    URIT-50 dan olingan ma'lumotlar:
    - blank.docx shablonidan yangi fayl yaratadi
    - {{SAMPLE}}, {{FISH}}, {{YOSH}} ni almashtiradi
    - {{NOMERATSIYA}} - ID ning oxirgi 2 raqami
    - Kod bo'yicha jadvaldagi Natija ustunini to'ldiradi
    - * bo'lgan ko'rsatkichlarni qizil (*) bilan chiqaradi
    - ✅ ID bo'lmasa ham, bemor topilmasa ham blanka HAR DOIM yaratiladi
    """
    # ✅ GLOBAL TRY-EXCEPT - HAR QANDAY XATO BO'LSA HAM BLANKA YARATILADI
    try:
        print("\n--- create_doc_from_data chaqirildi ---")
        print(f"✅ ID topilmasa ham blanka HAR DOIM yaratiladi")
        print(f"   patient_data: {patient_data}")
        print(f"   fish: '{fish}', yosh: '{yosh}', tugilgan_sana: '{tugilgan_sana}'")
        print(f"TEMPLATE_PATH: {TEMPLATE_PATH}")
        print(f"Shablon fayl mavjudmi: {os.path.exists(TEMPLATE_PATH)}")

        if not os.path.exists(TEMPLATE_PATH):
            print(f"❌ CRITICAL ERROR: Shablon topilmadi: {TEMPLATE_PATH}")
            print(f"   Joriy papka: {os.getcwd()}")
            print(f"   Script papkasi: {SCRIPT_DIR}")
            print(f"   ⚠️ URIT 50 papkasida 'blank.docx' fayli bo'lishi kerak!")
            print(f"   Tekshiring: {os.path.join(SCRIPT_DIR, 'blank.docx')}")
            # ❗ Shablon topilmasa blanka yaratib bo'lmaydi - bu kritik xato
            # Lekin dastur davom etishi kerak
            raise FileNotFoundError(f"Shablon topilmadi: {TEMPLATE_PATH}")

        try:
            print(f"Shablonni ochayapman: {TEMPLATE_PATH}")
            doc = Document(TEMPLATE_PATH)
            print("OK: Shablon ochildi")
        except Exception as e:
            print(f"❌ CRITICAL ERROR: Shablonni ochishda xato: {e}")
            import traceback
            traceback.print_exc()
            # ❗ Shablon ochib bo'lmasa blanka yaratib bo'lmaydi - bu kritik xato
            raise

        sample_no = (data.get("NO") or "").strip()

        # ✅ sample_id ning oxirgi 3 raqamini Nomeratsiya uchun olish (barcode skaner qilinganda)
        # Agar barcode skaner qilinmagan bo'lsa, sample NO dan olinadi
        nomeratsiya = ""
        print(f"\n🔢 Nomeratsiya (ID raqam) aniqlash:")
        print(f"   patient_data: {patient_data is not None}")
        
        if patient_data:
            # Avval sample_id dan olish (to'g'ri ID - masalan: 260220000919)
            sample_id = patient_data.get('sample_id', '')
            if sample_id:
                sample_id_str = str(sample_id).strip()
                print(f"   Sample ID topildi: {sample_id_str}")
                # sample_id ning oxirgi 3 raqami (barcode skaner qilinganda)
                if len(sample_id_str) >= 3:
                    nomeratsiya = sample_id_str[-3:]
                elif len(sample_id_str) == 2:
                    nomeratsiya = "0" + sample_id_str
                elif len(sample_id_str) == 1:
                    nomeratsiya = "00" + sample_id_str
            else:
                print(f"   Sample ID topilmadi, bemor ID dan olinmoqda")
                # Agar sample_id bo'lmasa, bemor ID dan olish (fallback)
                bemor_id = str(patient_data.get('id', ''))
                if bemor_id:
                    print(f"   Bemor ID topildi: {bemor_id}")
                    if len(bemor_id) >= 3:
                        nomeratsiya = bemor_id[-3:]
                    elif len(bemor_id) == 2:
                        nomeratsiya = "0" + bemor_id
                    elif len(bemor_id) == 1:
                        nomeratsiya = "00" + bemor_id
        else:
            print(f"   patient_data topilmadi")
        
        # Barcode skaner qilinmagan - sample NO dan olamiz
        if not nomeratsiya and sample_no:
            print(f"   Sample NO orqali harakat qilmoqda: {sample_no}")
            if len(sample_no) >= 3:
                nomeratsiya = sample_no[-3:]
            elif len(sample_no) == 2:
                nomeratsiya = "0" + sample_no
            elif len(sample_no) == 1:
                nomeratsiya = "00" + sample_no
        
        # Agar hali ham bo'sh bo'lsa, standart "000" yozamiz
        if not nomeratsiya:
            nomeratsiya = "000"
            print(f"   Nomeratsiya bo'sh, standart: {nomeratsiya}")
        
        print(f"✅ Yakuniy nomeratsiya: '{nomeratsiya}'")

        # ✅ Test vaqti HAR DOIM kompyuterdan olinadi
        # URIT-50 dan kelgan vaqt e'tiborsiz qoldiriladi (analizator vaqti ochib yonganda xato bo'ladi)
        # Har safar yangi vaqt olinadi
        dt_str = datetime.now().strftime("%d.%m.%Y %H:%M")
        print(f"\n🕐 Test vaqti (kompyuterdan, har safar yangi): {dt_str}")
        print(f"   ⚠️ URIT-50 dan kelgan vaqt e'tiborsiz qoldirildi (analizator vaqti xato bo'lishi mumkin)")

        # ✅ № o'rniga to'liq ID raqami ko'rsatiladi (barcha analizatorlar uchun bir xil)
        # Foydalanuvchi talabi: № o'rniga to'liq 12 xonalik ID raqami (masalan: 260111000030)
        full_id_str = ""
        if patient_data and patient_data.get('id'):
            full_id_str = str(patient_data.get('id'))
        else:
            full_id_str = "000000"
        bemor_id_str = full_id_str  # № o'rniga to'liq ID raqami
        print(f"   ✅ № o'rniga to'liq ID raqami ko'rsatiladi: '{bemor_id_str}'")
        
        # Oddiy placeholderlar - barcha joylarda qidirish
        print(f"\n📝 Placeholderlarni almashtiryapman:")
        print(f"   {{SAMPLE}}: '{sample_no}'")
        print(f"   {{DATE_TIME}}: '{dt_str}'")
        print(f"   {{FISH}}: '{fish}'")
        print(f"   {{YOSH}}: '{yosh}'")
        print(f"   {{NOMERATSIYA}}: '{nomeratsiya}'")
        print(f"   Bemor ID (№ o'rniga): '{bemor_id_str}'")
        
        # Placeholderlarni almashtirish - barcha variantlarni qo'llab-quvvatlash
        replace_text_in_doc(doc, "{{SAMPLE}}", sample_no or "")
        replace_text_in_doc(doc, "{{DATE_TIME}}", dt_str)
        replace_text_in_doc(doc, "{{DATE}}", dt_str.split()[0] if " " in dt_str else dt_str)  # Faqat sana
        replace_text_in_doc(doc, "{{TIME}}", dt_str.split()[1] if " " in dt_str else "")  # Faqat vaqt
        
        # FISH - FAQAT ID TOPILSA bemor ma'lumotlari yoziladi
        # ID topilmasa bemor ma'lumotlari bo'sh qoladi
        if patient_data and patient_data.get('id'):
            # ID topilgan - bemor ma'lumotlarini yozamiz
            if not fish and patient_data:
                fish = patient_data.get('fish', '') or patient_data.get('kod_yollanma', '') or patient_data.get('natija_kodi', '') or ''
            fish_value = (fish or '').strip()
            yosh_value = (yosh or '').strip()
        else:
            # ID topilmagan - bemor ma'lumotlari bo'sh qoladi
            fish_value = ""  # yoki "—" agar kerak bo'lsa
            yosh_value = ""
            print(f"   ⚠️ ID topilmadi - bemor ma'lumotlari bo'sh qoldirildi")
        
        replace_text_in_doc(doc, "{{FISH}}", fish_value)
        print(f"   ✅ {{FISH}} almashtirildi: '{fish_value}'")
        # fish ni fayl nomi uchun ham yangilaymiz
        fish = fish_value
        
        replace_text_in_doc(doc, "{{YOSH}}", yosh_value)
        
        replace_text_in_doc(doc, "{{NOMERATSIYA}}", nomeratsiya)
        # Boshqa variantlar ham bo'lishi mumkin
        replace_text_in_doc(doc, "{{NOMER}}", nomeratsiya)
        
        # ID - to'liq ID raqam (patient_data dan)
        replace_text_in_doc(doc, "{{ID}}", bemor_id_str)  # To'liq ID raqam
        replace_text_in_doc(doc, "{{№}}", bemor_id_str)  # № o'rniga to'liq ID raqami
        print(f"   ✅ {{ID}} almashtirildi: '{bemor_id_str}' (to'liq ID)")
        print(f"   ✅ № almashtirildi: '{bemor_id_str}' (to'liq ID raqami)")
        
        # № bilan boshlanuvchi raqamlarni ID bilan almashtirish
        # Formatni buzmasdan, replace_text_in_doc funksiyasidan foydalanamiz
        import re
        # Avval № bilan boshlanuvchi raqamlarni topamiz va ularni placeholderga aylantiramiz
        # Keyin placeholderlarni almashtiramiz
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text:
                            # № bilan boshlanuvchi raqamlarni topish va almashtirish
                            # Formatni buzmasdan replace_text_in_doc funksiyasidan foydalanamiz
                            # № bilan boshlanuvchi raqamlarni topish va ID bilan almashtirish
                            # Pattern: "№ 000008" yoki "№000008" yoki "000008" (№ belgisiz ham)
                            pattern = r'(№\s*)\d+|^(\d{6,})'  # № bilan yoki 6+ raqam
                            if re.search(pattern, para.text):
                                # Har bir run ichida almashtirish
                                for run in para.runs:
                                    run_text = run.text
                                    if re.search(pattern, run_text):
                                        # Formatni saqlab qolish
                                        original_font = run.font.name
                                        original_size = run.font.size
                                        original_bold = run.bold
                                        original_color = run.font.color.rgb if run.font.color.rgb else None
                                        
                                        # Matnni almashtirish
                                        # № bilan boshlanganlarni almashtirish
                                        if '№' in run_text:
                                            new_text = re.sub(r'(№\s*)\d+', r'\1' + bemor_id_str, run_text)
                                        else:
                                            # № belgisiz raqamlarni almashtirish (masalan: "000008")
                                            new_text = re.sub(r'\b\d{6,}\b', bemor_id_str, run_text)
                                        
                                        run.text = new_text
                                        
                                        # Formatni qayta tiklash
                                        if original_font:
                                            run.font.name = original_font
                                        if original_size:
                                            run.font.size = original_size
                                        if original_bold is not None:
                                            run.bold = original_bold
                                        if original_color:
                                            run.font.color.rgb = original_color
                                        print(f"   ✅ № raqam almashtirildi (jadval): '{para.text}' -> '{new_text}'")
                                        break
        
        # Oddiy paragraphlarda ham № ni almashtirish
        for para in doc.paragraphs:
            if para.text:
                # № bilan boshlanuvchi raqamlarni topish va ID bilan almashtirish
                pattern = r'(№\s*)\d+|^(\d{6,})'  # № bilan yoki 6+ raqam
                if re.search(pattern, para.text):
                    # Har bir run ichida almashtirish
                    for run in para.runs:
                        run_text = run.text
                        if re.search(pattern, run_text):
                            # Formatni saqlab qolish
                            original_font = run.font.name
                            original_size = run.font.size
                            original_bold = run.bold
                            original_color = run.font.color.rgb if run.font.color.rgb else None
                            
                            # Matnni almashtirish
                            # № bilan boshlanganlarni almashtirish
                            if '№' in run_text:
                                new_text = re.sub(r'(№\s*)\d+', r'\1' + bemor_id_str, run_text)
                            else:
                                # № belgisiz raqamlarni almashtirish (masalan: "000008")
                                new_text = re.sub(r'\b\d{6,}\b', bemor_id_str, run_text)
                            
                            run.text = new_text
                            
                            # Formatni qayta tiklash
                            if original_font:
                                run.font.name = original_font
                            if original_size:
                                run.font.size = original_size
                            if original_bold is not None:
                                run.bold = original_bold
                            if original_color:
                                run.font.color.rgb = original_color
                            print(f"   ✅ № raqam almashtirildi (paragraph): '{para.text}' -> '{new_text}'")
                            break
        
        # Vaqtni boshqa formatlarda ham qidirish (agar bo'lsa)
        date_only = dt_str.split()[0] if " " in dt_str else dt_str
        time_only = dt_str.split()[1] if " " in dt_str else ""
        replace_text_in_doc(doc, "{{DATE}}", date_only)  # Faqat sana
        replace_text_in_doc(doc, "{{TIME}}", time_only)  # Faqat vaqt
        print(f"✅ Placeholderlar almashtirildi")
        
        # Tug'ilgan sana - faqat yilni olish (ikkinchi rasmdagidek)
        tugilgan_yil = ""
        if tugilgan_sana:
            try:
                # Tug'ilgan sanani tozalash va formatlash
                tugilgan_sana_clean = str(tugilgan_sana).strip()
                
                if "-" in tugilgan_sana_clean:
                    parts = tugilgan_sana_clean.split("-")
                    if len(parts) >= 1:
                        tugilgan_yil = parts[0]  # Faqat yil (masalan: 2025-01-01 -> 2025)
                    # To'liq sana formatlash (DD.MM.YYYY)
                    if len(parts) == 3:
                        tugilgan_formatted = f"{parts[2]}.{parts[1]}.{parts[0]}"  # 2025-01-01 -> 01.01.2025
                    else:
                        tugilgan_formatted = tugilgan_sana_clean
                else:
                    # Agar boshqa format bo'lsa, yilni ajratishga harakat qilamiz
                    # Masalan: "1990" yoki "1990-12-25"
                    tugilgan_yil = tugilgan_sana_clean[:4] if len(tugilgan_sana_clean) >= 4 else ""
                    tugilgan_formatted = tugilgan_sana_clean
                
                print(f"   Tug'ilgan yil: '{tugilgan_yil}' (sana: '{tugilgan_sana_clean}')")
                replace_text_in_doc(doc, "{{TUGILGAN_SANA}}", tugilgan_formatted)
                replace_text_in_doc(doc, "{{TUGILGAN_YIL}}", tugilgan_yil)
            except Exception as e:
                print(f"   ⚠️ Tug'ilgan sana formatlashda xato: {e}")
                replace_text_in_doc(doc, "{{TUGILGAN_SANA}}", str(tugilgan_sana))
                # Yilni ajratishga harakat qilish
                try:
                    tugilgan_yil = str(tugilgan_sana)[:4] if len(str(tugilgan_sana)) >= 4 else ""
                    replace_text_in_doc(doc, "{{TUGILGAN_YIL}}", tugilgan_yil)
                except:
                    replace_text_in_doc(doc, "{{TUGILGAN_YIL}}", "")
        else:
            print(f"   ⚠️ Tug'ilgan sana topilmadi")
            replace_text_in_doc(doc, "{{TUGILGAN_SANA}}", "")
            replace_text_in_doc(doc, "{{TUGILGAN_YIL}}", "")

        # * bo'lgan kodlar
        abnormal_codes = set(data.get("ABNORMAL", []))

        # ===== JADVAL: Kod → Natija (3–4 ustun) =====
        # Jadval tuzilmasi:
        # 0 = Ko'rsatkich
        # 1,2 = Kod (birlashtirilgan)
        # 3,4 = Natija (birlashtirilgan)
        # 5,6 = Norma (birlashtirilgan)
        # 7   = O'lchov birligi
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 5:
                    continue

                # ✅ Sarlavha qatorini aniqlash - agar sarlavha bo'lsa, o'tkazib yuboramiz
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text + " "
                row_text_upper = row_text.upper()
                
                # Sarlavha qatorlarni aniqlash
                is_header_row = any(keyword in row_text_upper for keyword in [
                    "FIZIK", "XOSSALARI", "KIMYOVIY", "MIKROSKOPIYASI", "CHO'KMASI",
                    "KO'RSATKICH", "KOD", "NATIJA", "NORMA", "O'LCHOV BIRLIGI",
                    "SIYDIK", "TAXLILI"
                ])
                
                # Agar sarlavha qatori bo'lsa, o'tkazib yuboramiz - saqlab qolamiz
                if is_header_row:
                    print(f"   📋 Sarlavha qatori topildi va saqlandi: '{row_text.strip()}'")
                    continue

                # Kod hujayrasini topamiz (1 yoki 2-ustun)
                code1 = row.cells[1].text.strip()
                code2 = row.cells[2].text.strip() if len(row.cells) > 2 else ""
                code = code1 or code2

                # ✅ Kod ustunini tozalash - agar "10-15" yoki boshqa noto'g'ri qiymat bo'lsa, tozalash
                # LEU kod ustunida "10-15" yozilgan bo'lishi mumkin, uni tozalash kerak
                # Avval Ko'rsatkich ustunida "Leykotsitlar" borligini tekshirish
                indicator_text = row.cells[0].text.strip().upper() if len(row.cells) > 0 else ""
                
                # Agar Ko'rsatkich ustunida "Leykotsitlar" bo'lsa va kod ustunida raqam bo'lsa, bu LEU qatori
                if "LEYKOTSITLAR" in indicator_text:
                    # Bu LEU qatori, kod ustunini tekshirish
                    if code1 and not code1.upper() in ["LEU"] and (any(c.isdigit() for c in code1) or "10-15" in code1 or "-" in code1):
                        # Kod ustunida "10-15" yoki boshqa raqam bo'lsa, tozalash - faqat "LEU" ni qoldirish
                        for para in row.cells[1].paragraphs:
                            for run in para.runs:
                                run.text = ""
                            if len(para.runs) == 0:
                                para.add_run("")
                            from docx.oxml.ns import qn
                            run = row.cells[1].paragraphs[0].add_run("LEU")
                            run.font.name = "Times New Roman"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                            code1 = "LEU"
                            print(f"   ✅ LEU kod ustuni tozalandi: '{code1}' -> 'LEU' (FIZIK-KIMYOVIY jadvalida)")
                    if code2 and not code2.upper() in ["LEU"] and (any(c.isdigit() for c in code2) or "10-15" in code2):
                        # Ustun 2 ni tozalash
                        for para in row.cells[2].paragraphs:
                            for run in para.runs:
                                run.text = ""
                            if len(para.runs) == 0:
                                para.add_run("")
                            code2 = ""
                    # Kodni to'g'ri o'rnatish
                    if code1.upper() == "LEU":
                        code = "LEU"
                    elif not code or code not in ANALYTES:
                        code = "LEU"  # Agar kod topilmasa, LEU bo'lsin

                if code not in ANALYTES:
                    continue

                value = data.get(code)
                if not value:
                    continue

                # Natija hujayralari (3 va 4-ustun) – ikkalasini ham to'ldiramiz
                result_cells = []
                if len(row.cells) > 3:
                    result_cells.append(row.cells[3])
                if len(row.cells) > 4:
                    result_cells.append(row.cells[4])

                for cell in result_cells:
                    # ✅ Formatni buzmasdan tozalash - faqat matnni o'chiramiz
                    # Katak strukturini saqlab qolish uchun paragraphlarni tozalaymiz
                    for para in cell.paragraphs:
                        # Barcha runlarni o'chiramiz, lekin paragraph strukturini saqlaymiz
                        for run in para.runs:
                            run.text = ""
                        
                        # Agar paragraph bo'sh bo'lsa, birinchi paragraphni tozalaymiz
                        if len(para.runs) == 0:
                            para.add_run("")
                    
                    # Birinchi paragraphni olamiz
                    para = cell.paragraphs[0]
                    
                    # Qatorlar orasi qisqa qilish (bir vaqtda ko'p ma'lumot sig'ishi uchun)
                    pf = para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0

                    # ✅ * belgisini tozalash - blankaga chiqmasligi kerak
                    # Value dan * belgisini olib tashlash
                    import re
                    cleaned_value = re.sub(r'^\*\s*', '', str(value))  # Boshidagi * belgisini olib tashlash
                    cleaned_value = re.sub(r'\*\s*', '', cleaned_value)  # Barcha * belgilarini olib tashlash
                    cleaned_value = cleaned_value.strip()
                    
                    # Agar LEU bo'lsa, qiymatni tozalash - faqat son va birlikni saqlash
                    # Masalan: "*+- 15 CELL/µL" -> "15 CELL/µL"
                    if code == "LEU":
                        # * belgisini va +- belgisini tozalash
                        cleaned_value = re.sub(r'^\*\s*[+\-]+\s*', '', cleaned_value)
                        cleaned_value = cleaned_value.strip()
                    
                    text_value = cleaned_value

                    # Birinchi runni topamiz yoki yangi run qo'shamiz
                    if len(para.runs) > 0:
                        run = para.runs[0]
                        run.text = text_value
                    else:
                        run = para.add_run(text_value)

                    # FORMAT: Times New Roman, 12, bold
                    from docx.oxml.ns import qn
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
                    run.font.size = Pt(12)
                    run.bold = True

                    # Abnormal kodlar qizil rangda, lekin * belgisi yo'q
                    if code in abnormal_codes:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # ====== MIKROSKOPIYA AVTOMATIK TO'LDIRISH ======
        print("\n" + "="*60)
        print("MIKROSKOPIYA qismini to'ldirish...")
        print("="*60)
        
        # LEU va BLD qiymatlarini olish
        leu_value = data.get("LEU", "0 CELL/µL")
        bld_value = data.get("BLD", "0 CELL/µL")
        
        print(f"LEU (biokimyo): {leu_value}")
        print(f"BLD (biokimyo): {bld_value}")
        
        # Mikroskopiya qiymatlarini hisoblash
        micro_values = get_microscopy_values(leu_value, bld_value)
        
        print("\nMikroskopiya natijalari:")
        for code, info in micro_values.items():
            color = "QIZIL" if info["is_red"] else "QORA"
            print(f"  {code}: {info['value']} ({color})")
        
        # ====== FAQAT MIKROSKOPIYA JADVALIGA YOZISH ======
        # "SIYDIK CHO'KMASI MIKROSKOPIYASI" sarlavhasidan keyin bo'lgan qatorlar
        
        microscopy_started = False
        
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                
                # Mikroskopiya qismini aniqlash
                if len(cells) > 0:
                    first_cell = cells[0].text.strip()
                    
                    # Agar "MIKROSKOPIYASI" so'zi bor bo'lsa, mikroskopiya qismi boshlandi
                    if "MIKROSKOPIYASI" in first_cell.upper() or "MIKROSKOPIYA" in first_cell.upper():
                        microscopy_started = True
                        print(f"✅ Mikroskopiya qismi topildi")
                        continue
                    
                    # Agar mikroskopiya qismi boshlanmagan bo'lsa, o'tkazib yuborish
                    if not microscopy_started:
                        continue
                
                # Mikroskopiya qismida bo'lsak, qatorlarni tekshirish
                if len(cells) >= 4:
                    cell_0_text = cells[0].text.strip()
                    cell_0_text_upper = cell_0_text.upper()
                    
                    # Sarlavha qatorlarini o'tkazib yuborish
                    if any(keyword in cell_0_text_upper for keyword in ["KO'RSATKICH", "KOD", "NATIJA", "NORMA", "MIKROSKOPIYASI", "CHO'KMASI"]):
                        continue
                    
                    # ====== LEYKOTSITLAR (WBC) ======
                    if "Leykotsitlar" in cell_0_text or "Лейкоцитлар" in cell_0_text or "LEYKOTSITLAR" in cell_0_text_upper:
                        # Kod ustuni tekshirish
                        cell_1_text = cells[1].text.strip() if len(cells) > 1 else ""
                        if "WBC" in cell_1_text.upper() or cell_1_text == "":
                            # Bu mikroskopiya qismi (WBC)
                            if "WBC" in micro_values:
                                value_info = micro_values["WBC"]
                                
                                # Natija ustunlariga yozish (2 va 3-ustunlar)
                                for col_idx in [2, 3]:
                                    if len(cells) > col_idx:
                                        cells[col_idx].paragraphs[0].clear()
                                        run = cells[col_idx].paragraphs[0].add_run(value_info["value"])
                                        run.font.name = "Times New Roman"
                                        run.font.size = Pt(12)
                                        if value_info["is_red"]:
                                            run.font.color.rgb = RGBColor(255, 0, 0)
                                            run.bold = True
                                        else:
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                                            run.bold = True
                                
                                print(f"  ✅ WBC: {value_info['value']} yozildi")
                    
                    # ====== ERITROSITLAR O'ZGARGAN (RBC-D) ======
                    # Faqat KOD ustunida qidirish - blankada "RBC-D" shaklida yozilgan
                    elif len(cells) > 1:
                        cell_1_text = cells[1].text.strip().upper()
                        # RBC-D kodini tekshirish
                        if "RBC-D" in cell_1_text:
                            # RBC-D ni tekshirish
                            if "RBC-D" in micro_values:
                                value_info = micro_values["RBC-D"]
                                
                                for col_idx in [2, 3]:
                                    if len(cells) > col_idx:
                                        cells[col_idx].paragraphs[0].clear()
                                        run = cells[col_idx].paragraphs[0].add_run(value_info["value"])
                                        run.font.name = "Times New Roman"
                                        run.font.size = Pt(12)
                                        if value_info["is_red"]:
                                            run.font.color.rgb = RGBColor(255, 0, 0)
                                            run.bold = True
                                        else:
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                                            run.bold = True
                                
                                print(f"  ✅ RBC-D: {value_info['value']} yozildi")
                        
                        # ====== ERITROSITLAR O'ZGARMAGAN (RBC-U) ======
                        # RBC-U kodini tekshirish - blankada "RBC-U" shaklida yozilgan
                        elif "RBC-U" in cell_1_text:
                            # RBC-U ni tekshirish
                            if "RBC-U" in micro_values:
                                value_info = micro_values["RBC-U"]
                                
                                for col_idx in [2, 3]:
                                    if len(cells) > col_idx:
                                        cells[col_idx].paragraphs[0].clear()
                                        run = cells[col_idx].paragraphs[0].add_run(value_info["value"])
                                        run.font.name = "Times New Roman"
                                        run.font.size = Pt(12)
                                        if value_info["is_red"]:
                                            run.font.color.rgb = RGBColor(255, 0, 0)
                                            run.bold = True
                                        else:
                                            run.font.color.rgb = RGBColor(0, 0, 0)
                                            run.bold = True
                                
                                print(f"  ✅ RBC-U: {value_info['value']} yozildi")
        
        print("="*60)
        print("✅ Mikroskopiya to'ldirildi!")
        print("="*60)

        # ✅ MIKROSKOPIYA QISMINI AVTOMATIK TO'LDIRISH (2-rasmdagi jadvalga qarab)
        fill_microscopy_results(doc, data, abnormal_codes)

        # ✅ Blankaning tepasidagi bo'sh joylarni qisqartirish va jadval balandligini o'rnatish
        # Foydalanuvchi talabi: tepada ko'p joy qolmasligi kerak
        # BU FUNKSIYA Sarlavhalarni saqlab qoladi
        optimize_page_spacing(doc)  # Tepadan 0.6 sm, jadval qatorlari 0.4 sm

        # ✅ Yakuniy format: hammasi Times New Roman 12, 'Test vaqti' 13
        # ✅ Faqat natijalar qismi formatlanadi (blankaning tepasi o'zgarmaydi)
        # Sarlavhalarni saqlab qolish uchun format_results_only ni chaqiramiz
        # IMPORTANT: optimize_page_spacing dan KEYIN chaqiriladi, chunki sarlavhalarni saqlab qolish kerak
        format_results_only(doc)

        # ✅ Test vaqti 13 bo'lib qoladi
        set_test_time_font_13(doc)


        # Sana bo'yicha papka
        save_dir = get_save_dir()
        print(f"\nSaqlash papkasi: {save_dir}")
        try:
            os.makedirs(save_dir, exist_ok=True)
            print(f"OK: Papka yaratildi yoki mavjud: {save_dir}")
        except Exception as e:
            print(f"⚠️ WARNING: Papka yaratishda xato: {e}")
            print(f"   Standart papkadan foydalanamiz: {BASE_DIR}")
            # ❗ Papka yaratib bo'lmasa ham blanka yaratilishi kerak
            save_dir = BASE_DIR
            try:
                os.makedirs(save_dir, exist_ok=True)
            except:
                pass  # Agar standart papka ham yaratib bo'lmasa, davom etamiz

        # ✅ Fayl nomini bemor ism-familiyasi bilan yaratish
        # Format: "SIYDIK TAHLILI - 260111000029 Shaydullayeva Oysha 2021.docx"
        def clean_filename(text: str) -> str:
            """Fayl nomida maxsus belgilarni tozalash"""
            # Maxsus belgilar: / \ : * ? " < > |
            forbidden_chars = r'[\/\\:*?"<>|]'
            cleaned = re.sub(forbidden_chars, '_', text)
            # Qo'shimcha bo'shliqlarni bitta bo'shliq qilish
            cleaned = re.sub(r'\s+', ' ', cleaned).strip()
            return cleaned
        
        # Bemor ma'lumotlarini olish (fayl nomi uchun)
        # ✅ ID None bo'lsa ham xato chiqmasligi uchun
        sample_id_raw = data.get("ID")
        sample_id_str = (sample_id_raw or "").strip() if sample_id_raw else ""
        if not sample_id_str and patient_data:
            sample_id_str = str(patient_data.get('sample_id', '') or '').strip()
        
        bemor_ismi = fish.strip() if fish else ""
        if not bemor_ismi and patient_data:
            bemor_ismi = (patient_data.get('fish', '') or 
                          patient_data.get('kod_yollanma', '') or 
                          patient_data.get('natija_kodi', '') or '').strip()
        
        # Tug'ilgan yilni olish (fayl nomi uchun)
        tugilgan_yil_str = ""
        if tugilgan_sana:
            try:
                tugilgan_sana_clean = str(tugilgan_sana).strip()
                if "-" in tugilgan_sana_clean:
                    tugilgan_yil_str = tugilgan_sana_clean.split("-")[0]
                else:
                    tugilgan_yil_str = tugilgan_sana_clean[:4] if len(tugilgan_sana_clean) >= 4 else ""
            except:
                pass
        
        # Fayl nomini yaratish
        if sample_id_str and bemor_ismi and tugilgan_yil_str:
            # To'liq format: "SIYDIK TAHLILI - 260111000029 Shaydullayeva Oysha 2021.docx"
            fname = f"SIYDIK TAHLILI - {sample_id_str} {clean_filename(bemor_ismi)} {tugilgan_yil_str}.docx"
        elif sample_id_str and bemor_ismi:
            # ID va ism bo'lsa
            fname = f"SIYDIK TAHLILI - {sample_id_str} {clean_filename(bemor_ismi)}.docx"
        elif sample_id_str:
            # Faqat ID bo'lsa
            fname = f"SIYDIK TAHLILI - {sample_id_str}.docx"
        elif sample_no:
            # Sample NO bo'lsa (eski format - fallback)
            fname = f"SIYDIK TAHLILI - NO_{sample_no}.docx"
        else:
            # Hech narsa bo'lmasa
            now_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            fname = f"SIYDIK TAHLILI - {now_str}.docx"

        full_path = os.path.join(save_dir, fname)

        try:
            print(f"\nBlanka faylni saqlayapman: {full_path}")
            print(f"   Papka: {save_dir}")
            print(f"   Fayl nomi: {fname}")
            doc.save(full_path)
            print(f"✅ OK: Word blank saqlandi: {full_path}")
            print(f"   Fayl mavjudmi: {os.path.exists(full_path)}")
            if os.path.exists(full_path):
                file_size = os.path.getsize(full_path)
                print(f"   Fayl o'lchami: {file_size} bayt")
            return full_path
        except Exception as e:
            print(f"❌ CRITICAL ERROR: Wordni saqlashda xato: {e}")
            import traceback
            traceback.print_exc()
            print(f"   Saqlash papkasi: {save_dir}")
            print(f"   Fayl nomi: {fname}")
            print(f"   To'liq yo'l: {full_path}")
            # ❗ Fayl saqlab bo'lmasa blanka yaratib bo'lmaydi - bu kritik xato
            raise
    except FileNotFoundError as file_err:
        # ❗ Shablon topilmasa - bu kritik xato, blanka yaratib bo'lmaydi
        print(f"\n❌ CRITICAL ERROR: Shablon topilmadi - blanka yaratib bo'lmaydi!")
        print(f"   Xato: {file_err}")
        print(f"   ⚠️ Dastur davom etadi, lekin blanka yaratilmadi")
        # Shablon topilmasa blanka yaratib bo'lmaydi - bu kritik xato
        return ""
    except Exception as outer_e:
        # ✅ BOSHQA XATOLARNI TUTISH - HAR QANDAY XATO BO'LSA HAM BLANKA YARATILISHIGA HARAKAT QILAMIZ
        print(f"\n❌ CRITICAL ERROR: create_doc_from_data da xato: {outer_e}")
        import traceback
        traceback.print_exc()
        print(f"   ⚠️ Blanka yaratilmadi, lekin dastur davom etadi")
        # ❗ Boshqa xatolar uchun ham blanka yaratib bo'lmaydi
        return ""



def read_one_block() -> str:
    """
    URIT-50 dan bitta natijani (0x03 / ETX gacha) o'qib qaytaradi.
    Har chaqirilganda COM4 portni ochib-yopadi.
    """
    try:
        with serial.Serial(
            port=COM_PORT,
            baudrate=BAUDRATE,
            bytesize=serial.EIGHTBITS,
            parity=serial.PARITY_NONE,
            stopbits=serial.STOPBITS_ONE,
            timeout=1,
        ) as ser:
            buffer = ""

            while True:
                chunk = ser.read(ser.in_waiting or 1)
                if not chunk:
                    continue

                text = chunk.decode(errors="ignore")
                buffer += text

                if "\x03" in buffer:
                    block, _, _ = buffer.partition("\x03")
                    return block
    except serial.SerialException as e:
        error_msg = f"COM port ({COM_PORT}) ochib bo'lmadi: {e}"
        print(f"⚠️ {error_msg}")
        print(f"   Tekshiring:")
        print(f"   1. URIT-50 qurilmasi ulanganmi?")
        print(f"   2. COM{COM_PORT[-1]} port mavjudmi? (Device Manager)")
        print(f"   3. Port boshqa dastur tomonidan ishlatilmoqdami?")
        print(f"   4. Port nomi to'g'rimi? (Hozirgi: {COM_PORT})")
        # Bo'sh string qaytarish - dastur to'xtamasligi uchun
        return ""
    except Exception as e:
        error_msg = f"COM port o'qishda xato: {e}"
        print(f"❌ {error_msg}")
        import traceback
        traceback.print_exc()
        return ""


# ====== ASOSIY SIKL: HAR TАHLIL UCHUN ALOHIDA ULANISH ======
def main():
    global current_barcode
    
    # ✅ Terminal encoding muammosini hal qilish
    import sys
    import io
    # Windows terminalda UTF-8 encoding ni ta'minlash
    if sys.platform == 'win32':
        try:
            sys.stdout.reconfigure(encoding='utf-8')
            sys.stderr.reconfigure(encoding='utf-8')
        except:
            # Agar reconfigure ishlamasa, encoding ni o'rnatish
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
            sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
    
    print("=" * 60)
    print("  URIT-50 Integratsiya - AzizMedLine LIMS")
    print("=" * 60)
    print("URIT-50 ma'lumot kutilyapti...")
    print("Skaner orqali barcode o'qib, bemor ma'lumotlari avtomatik olinadi")
    print("Natijalar blankaga avtomatik yoziladi\n")

    # Skaner thread ni ishga tushirish
    scanner_thread = threading.Thread(target=barcode_scanner_thread, daemon=True)
    scanner_thread.start()
    print("OK: Skaner thread ishga tushdi\n")

    count = 0

    while True:
        # 1) URIT-50 dan bitta natija blokini o'qiymiz
        print("\n" + "=" * 60)
        print("📡 Yangi tahlilni yuborish uchun URIT'da RS232 tugmasini bosing...")
        print("📷 Yoki avval barcode skanerlang (kod_yollanma yoki natija_kodi)")
        print("=" * 60)
        
        block = read_one_block()
        
        # Agar block bo'sh bo'lsa (port ochilmagan), kutish va qayta urinish
        if not block:
            print(f"\n⚠️ COM port ochilmadi yoki ma'lumot kelmadi. 5 soniyadan keyin qayta uriniladi...")
            time.sleep(5)
            continue  # Keyingi iteratsiyaga o'tish
        
        count += 1

        print(f"\n======= #{count}-TAHLIL RAW BLOK =======")
        print(block)
        print("=======================================")

        # 2) RAW logga saqlaymiz
        raw_path = save_raw_block(block)
        print("📝 RAW saqlandi:", raw_path)

        # 3) Parse qilamiz
        data = parse_urit_block(block)
        print("\n--- PARSE QILINGAN NATIJA ---")
        print(data)

        # 4) Bemor ma'lumotlarini bazadan olish
        patient_data = None
        fish = ""
        yosh = ""
        tugilgan_sana = ""
        
        # Avval queue dan yangi barcode ni olish (agar bo'lsa)
        try:
            while not barcode_queue.empty():
                new_barcode = barcode_queue.get_nowait()
                if new_barcode:
                    current_barcode = new_barcode
                    print(f"📷 Queue dan yangi barcode olingan: {current_barcode}")
        except queue.Empty:
            pass
        
        # Barcode ni aniqlash - prioritet tartibida
        barcode_to_search = None
        
        # 1) Avval current_barcode ni tekshiramiz (skaner orqali kelgan) - ENG YUQORI USTUNLIK
        if current_barcode:
            barcode_to_search = current_barcode
            print(f"📷 Skaner orqali kelgan barcode: {barcode_to_search}")
        
        # 2) Agar current_barcode bo'lmasa, FAQAT sample_id (ID) bilan qidirish
        if not barcode_to_search and data.get("ID"):
            sample_id_from_urit = data.get("ID").strip()
            if sample_id_from_urit and len(sample_id_from_urit) >= 10:  # 12 xonalik ID
                barcode_to_search = sample_id_from_urit
                print(f"🆔 URIT-50 dan kelgan sample_id: {barcode_to_search}")
            else:
                print(f"⚠️ ID juda qisqa yoki noto'g'ri: {sample_id_from_urit}")
        
        # 3) NO ni umuman ishlatmaslik (chunki har safar 000001 dan boshlanadi)
        # Agar ID topilmasa, bemor topilmaydi
        if not barcode_to_search:
            print(f"⚠️ OGOHLANTIRISH: Sample ID topilmadi!")
            print(f"   - Skaner bilan barcode o'qing, yoki")
            print(f"   - URIT-50 da to'g'ri sample ID kiriting")
        
        if barcode_to_search:
            print(f"\n🔍 Bemor qidiryapman - Barcode: '{barcode_to_search}'")
            patient_data = get_patient_by_barcode(barcode_to_search)
            
            if patient_data:
                # ✅ Fish bo'sh bo'lsa, kod_yollanma yoki boshqa maydonni ishlatish
                fish = patient_data.get('fish', '') or patient_data.get('kod_yollanma', '') or patient_data.get('natija_kodi', '') or ''
                yosh = str(patient_data.get('yosh', ''))
                
                # ✅ Tug'ilgan sanani olish - barcha variantlarni tekshirish
                tugilgan_sana = ""
                if patient_data.get('tugilgan_sana'):
                    tugilgan_sana = str(patient_data['tugilgan_sana'])
                    print(f"   ✅ Tug'ilgan sana topildi: '{tugilgan_sana}'")
                else:
                    print(f"   ⚠️ Tug'ilgan sana topilmadi (bazada NULL yoki bo'sh)")
                    tugilgan_sana = ""
                
                print(f"\n✅ Bemor ma'lumotlari topildi:")
                print(f"   F.I.SH: '{fish}' (fish bo'sh bo'lsa, kod_yollanma ishlatildi)")
                print(f"   Yosh: {yosh}")
                print(f"   Tug'ilgan sana: '{tugilgan_sana}'")
                print(f"   ID: {patient_data.get('id')}")
                print(f"   Kod yo'llanma: {patient_data.get('kod_yollanma')}")
                print(f"   Natija kodi: {patient_data.get('natija_kodi')}")
                print(f"   Sample ID: {patient_data.get('sample_id', 'NULL')}")
            else:
                print(f"\n⚠️ Bemor topilmadi: '{barcode_to_search}'")
                print("   Blanka bemor ma'lumotlarisiz yaratiladi.")
                print("   Tekshiring:")
                print("   - Barcode to'g'ri skanerlanganmi?")
                print("   - Bemor bazada mavjudmi? (kod_yollanma yoki natija_kodi)")
        else:
            print("\n⚠️ Barcode topilmadi!")
            print("   Qidiruv uchun quyidagilardan bittasini kiriting:")
            print("   1. Barcode skanerlang (kod_yollanma yoki natija_kodi)")
            print("   2. URIT-50 da to'g'ri Sample ID ni kiriting")
            print("   Blanka bemor ma'lumotlarisiz yaratiladi.")

        # 4b) Strip turini aniqlash va 11-param uchun ortiqcha maydonlarni tozalash
        strip_type = detect_strip_type(data)
        print(f"\n📋 Poloskа turi: {strip_type}-parametrli")
        if strip_type == 11:
            print("   → 11-parametrli: MA, Ca, CR, ACR blankada ko'rinmaydi")
            for _extra in ("MA", "Ca", "CR", "ACR"):
                data[_extra] = ""

        # 5) Blanka yaratamiz
        try:
            print("\n" + "=" * 60)
            print("Blanka yaratishni boshlayapman...")
            print("=" * 60)
            out_path = create_doc_from_data(data, fish=fish, yosh=yosh, tugilgan_sana=tugilgan_sana, patient_data=patient_data)
            if out_path:
                print(f"\n✅ OK: Blanka saqlandi:\n{out_path}\n")
                
                # 6) Bazaga saqlash - FAQAT ID TOPILSA
                if patient_data and patient_data.get('id'):
                    print("💾 Bazaga saqlayapman...")
                    try:
                        save_result_to_db(data.get("NO"), data, patient_data, out_path)
                        print("✅ Bazaga saqlandi")
                    except Exception as db_err:
                        print(f"⚠️ Bazaga saqlashda xato: {db_err}")
                        print("   Blanka yaratildi, lekin bazaga saqlanmadi")
                else:
                    print("ℹ️ ID topilmadi - bazaga saqlash o'tkazilmadi")
                    print("   Blanka yaratildi va faylga saqlandi")
                
                # Barcode ni tozalash (keyingi tahlil uchun yangi barcode skanerlash uchun)
                # Faqat agar bitta tahlil uchun ishlatilgan bo'lsa, tozalaymiz
                # Aks holda, keyingi tahlil uchun saqlab qolamiz
                print("\n" + "=" * 60)
                print("OK: Tahlil yakunlandi. Keyingi tahlil uchun barcode skanerlang...")
                print("=" * 60)
                # current_barcode ni tozalashni o'chirib qo'yamiz, chunki keyingi tahlil uchun ham kerak bo'lishi mumkin
                # current_barcode = None  # Faqat kerak bo'lganda tozalang
            else:
                print("\nWARNING: Blanka saqlanmadi (create_doc_from_data bo'sh string qaytardi).")
                print("   Muammo:")
                print("   - Shablon fayl topilmaganmi?")
                print("   - Faylni saqlashda xatolik bo'ldimi?")
                print("   - Boshqa xatolik yuz berdimimi?\n")
        except Exception as e:
            import traceback
            print("\nERROR: Blanka yaratishda xato:")
            print(f"   Xato turi: {type(e).__name__}")
            print(f"   Xato matni: {str(e)}")
            traceback.print_exc()
            print("Sikl davom etadi. Keyingi tahlil...\n")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️ Dastur to'xtatildi (Ctrl+C)")
        db_close()  # Database ulanishini yopish
    except Exception as e:
        print(f"\n\n❌ Xatolik: {e}")
        db_close()  # Database ulanishini yopish
        raise
    finally:
        db_close()  # Har doim database ulanishini yopish
