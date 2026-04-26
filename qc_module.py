# -*- coding: utf-8 -*-
"""
Sifat Nazorati (Quality Control) moduli — YENGIL VERSIYA
Laboratoriya analizatorlari uchun QC tizimi.
matplotlib ISHLATILMAYDI — faqat tk.Canvas (yengil, tez).
Barcha DB operatsiyalari fon thread da bajariladi (UI qotmaydi).
"""

import os
import json
import tkinter as tk
from tkinter import ttk, messagebox
from datetime import datetime, date, timedelta
import threading
import math

try:
    import mysql.connector
    from monoblok_db_config import DB_CONFIG
    DB_AVAILABLE = True
except ImportError:
    DB_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════
# GEMOTOLOGIYA STANDART PARAMETRLARI (19 ta)
# ══════════════════════════════════════════════════════════════════
HEMATOLOGY_PARAMS = [
    ("WBC",     "10^9/L"),
    ("LYM#",    "10^9/L"),
    ("LYM%",    "%"),
    ("MID#",    "10^9/L"),
    ("MID%",    "%"),
    ("GRAN#",   "10^9/L"),
    ("GRAN%",   "%"),
    ("RBC",     "10^12/L"),
    ("HGB",     "g/L"),
    ("HCT",     "%"),
    ("MCV",     "fL"),
    ("MCH",     "pg"),
    ("MCHC",    "g/L"),
    ("RDW-CV",  "%"),
    ("RDW-SD",  "fL"),
    ("PLT",     "10^9/L"),
    ("MPV",     "fL"),
    ("PDW",     "%"),
    ("PCT",     "%"),
]

# QC da kerak bo'lmagan testlar
QC_EXCLUDED_BIO = {'bilirubin', 'revmoproba', 'rf', 'aslo', 'crp', 'timol'}
QC_BILIRUBIN_REPLACE = [("Umumiy bilirubin", "mkmol/l"), ("Bog'langan bilirubin", "mkmol/l")]


# ══════════════════════════════════════════════════════════════════
# DB — har safar yangi ulanish (pool emas, qotmasligi uchun)
# ══════════════════════════════════════════════════════════════════
def _db_conn():
    if not DB_AVAILABLE:
        return None
    try:
        return mysql.connector.connect(**DB_CONFIG, connection_timeout=5)
    except Exception as e:
        print(f"[QC] DB xato: {e}")
        return None


def _run_in_bg(func, callback=None):
    """func ni fon thread da ishga tushiradi, callback(result) ni UI thread da chaqiradi."""
    def _worker():
        try:
            result = func()
        except Exception as e:
            print(f"[QC] BG xato: {e}")
            result = None
        if callback:
            try:
                # Tkinter after(0) — UI thread ga qaytish
                _root_ref[0].after(0, lambda: callback(result))
            except Exception:
                pass
    threading.Thread(target=_worker, daemon=True).start()

_root_ref = [None]  # UI thread uchun root reference


# ══════════════════════════════════════════════════════════════════
# WESTGARD QOIDALARI
# ══════════════════════════════════════════════════════════════════
def evaluate_westgard(values, target, sd):
    if not values or sd <= 0:
        return 'in_control', ''
    n = len(values)
    zs = [(v - target) / sd for v in values]
    last_z = zs[-1]
    if abs(last_z) > 3.0:
        return 'out_of_control', '1-3s'
    if abs(last_z) > 2.0:
        if n >= 2 and abs(zs[-2]) > 2.0:
            if (zs[-1] > 0 and zs[-2] > 0) or (zs[-1] < 0 and zs[-2] < 0):
                return 'out_of_control', '2-2s'
        if n >= 2 and abs(zs[-1] - zs[-2]) > 4.0:
            return 'out_of_control', 'R-4s'
        return 'warning', '1-2s'
    if n >= 4:
        last4 = zs[-4:]
        if all(z > 1.0 for z in last4) or all(z < -1.0 for z in last4):
            return 'out_of_control', '4-1s'
    if n >= 10:
        last10 = zs[-10:]
        if all(z > 0 for z in last10) or all(z < 0 for z in last10):
            return 'out_of_control', '10x'
    return 'in_control', ''


# ══════════════════════════════════════════════════════════════════
# DB JADVALLARINI YARATISH (fon thread da)
# ══════════════════════════════════════════════════════════════════
def _ensure_qc_tables_sync():
    conn = _db_conn()
    if not conn:
        return False
    try:
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS qc_lots (
                id INT AUTO_INCREMENT PRIMARY KEY,
                lot_number VARCHAR(50) NOT NULL,
                analyzer_type ENUM('hematology','biochemistry') NOT NULL,
                level VARCHAR(20) DEFAULT 'Level 1',
                expiry_date DATE,
                description TEXT,
                created_at DATETIME DEFAULT NOW(),
                UNIQUE KEY uk_lot (lot_number, analyzer_type, level)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS qc_lot_params (
                id INT AUTO_INCREMENT PRIMARY KEY,
                lot_id INT NOT NULL,
                parameter_name VARCHAR(100) NOT NULL,
                target_value DECIMAL(12,4),
                sd_value DECIMAL(12,4),
                unit VARCHAR(50),
                FOREIGN KEY (lot_id) REFERENCES qc_lots(id) ON DELETE CASCADE,
                UNIQUE KEY uk_param (lot_id, parameter_name)
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS qc_results (
                id INT AUTO_INCREMENT PRIMARY KEY,
                lot_param_id INT NOT NULL,
                measured_value DECIMAL(12,4) NOT NULL,
                measured_at DATETIME DEFAULT NOW(),
                source VARCHAR(50) DEFAULT 'manual',
                westgard_status VARCHAR(50) DEFAULT 'in_control',
                note TEXT,
                FOREIGN KEY (lot_param_id) REFERENCES qc_lot_params(id) ON DELETE CASCADE
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """)
        conn.commit()
        return True
    except Exception as e:
        print(f"[QC] Jadval xato: {e}")
        return False
    finally:
        conn.close()


def _load_bio_params_sync():
    conn = _db_conn()
    if not conn:
        return []
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT tahlil_nomi, birlik FROM tahlillar_norma WHERE guruh = 'BIO' ORDER BY tahlil_nomi")
        rows = cursor.fetchall()
        result = []
        bilirubin_added = False
        for r in rows:
            name_lower = r['tahlil_nomi'].lower()
            skip = False
            for excl in QC_EXCLUDED_BIO:
                if excl in name_lower:
                    if excl == 'bilirubin' and not bilirubin_added:
                        result.extend(QC_BILIRUBIN_REPLACE)
                        bilirubin_added = True
                    skip = True
                    break
            if not skip:
                result.append((r['tahlil_nomi'], r['birlik'] or ''))
        return result
    except Exception as e:
        print(f"[QC] BIO params xato: {e}")
        return []
    finally:
        conn.close()


# ══════════════════════════════════════════════════════════════════
# LEVEY-JENNINGS GRAFIK (tk.Canvas — matplotlib siz)
# ══════════════════════════════════════════════════════════════════
class LJChart:
    """Levey-Jennings grafik — sof tk.Canvas bilan"""

    def __init__(self, parent):
        self.canvas = tk.Canvas(parent, bg="white", highlightthickness=1,
                                highlightbackground="#CCCCCC")
        self.canvas.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        self.canvas.bind("<Configure>", lambda e: self._last_data and self.draw(*self._last_data))
        self._last_data = None

    def draw(self, values, target, sd, param_name=""):
        self._last_data = (values, target, sd, param_name)
        c = self.canvas
        c.delete("all")
        w = c.winfo_width()
        h = c.winfo_height()
        if w < 50 or h < 50:
            return

        pad_l, pad_r, pad_t, pad_b = 55, 15, 30, 25

        if not values or target <= 0 or sd <= 0:
            c.create_text(w // 2, h // 2, text="Ma'lumot yo'q", fill="#999", font=("Arial", 11))
            return

        plot_w = w - pad_l - pad_r
        plot_h = h - pad_t - pad_b
        n = len(values)

        y_min = target - 4 * sd
        y_max = target + 4 * sd

        def to_x(i):
            if n == 1:
                return pad_l + plot_w / 2
            return pad_l + i * plot_w / (n - 1)

        def to_y(v):
            if y_max == y_min:
                return pad_t + plot_h / 2
            return pad_t + (1 - (v - y_min) / (y_max - y_min)) * plot_h

        # Sarlavha
        c.create_text(w // 2, 12, text=f"{param_name} — Levey-Jennings",
                       font=("Arial", 10, "bold"), fill="#333")

        # SD chiziqlari
        sd_lines = [
            (target + 3*sd, "#FF0000", "+3SD"),
            (target + 2*sd, "#FF8C00", "+2SD"),
            (target + sd,   "#228B22", "+1SD"),
            (target,        "#000000", "X\u0304"),
            (target - sd,   "#228B22", "-1SD"),
            (target - 2*sd, "#FF8C00", "-2SD"),
            (target - 3*sd, "#FF0000", "-3SD"),
        ]
        for val, color, label in sd_lines:
            y = to_y(val)
            dash = (4, 3) if "SD" in label and "3" not in label else ()
            lw = 2 if "3SD" in label or label == "X\u0304" else 1
            c.create_line(pad_l, y, w - pad_r, y, fill=color, width=lw, dash=dash)
            c.create_text(pad_l - 5, y, text=label, anchor="e", fill=color, font=("Arial", 8))

        # Nuqtalar va chiziq
        points = []
        for i, v in enumerate(values):
            x = to_x(i)
            y = to_y(v)
            points.append((x, y, v))

        # Chiziq
        if len(points) > 1:
            coords = []
            for x, y, _ in points:
                coords.extend([x, y])
            c.create_line(*coords, fill="#4169E1", width=1)

        # Nuqtalar
        for i, (x, y, v) in enumerate(points):
            z = abs((v - target) / sd) if sd > 0 else 0
            color = "#FF0000" if z > 3 else "#FF8C00" if z > 2 else "#4169E1"
            r = 4
            c.create_oval(x-r, y-r, x+r, y+r, fill=color, outline=color)
            # Raqam yozish (kam bo'lsa)
            if n <= 31:
                c.create_text(x, pad_t + plot_h + 12, text=str(i+1), fill="#666", font=("Arial", 7))

    def clear(self):
        self._last_data = None
        self.canvas.delete("all")


# ══════════════════════════════════════════════════════════════════
# ASOSIY QC OYNASI
# ══════════════════════════════════════════════════════════════════
def open_qc_window(parent=None):
    """Sifat nazorati oynasini ochish — UI darhol ochiladi, DB fon da."""
    win = tk.Toplevel(parent)
    win.title("Sifat Nazorati — Quality Control")
    win.geometry("1300x800")
    win.configure(bg="#F5F5F5")
    _root_ref[0] = win

    # DB jadvallarni fon da yaratish
    _run_in_bg(_ensure_qc_tables_sync)

    current_lot_id = [None]
    current_param_id = [None]

    # ── Notebook ──────────────────────────────────────────────────
    notebook = ttk.Notebook(win)
    notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

    # ═══════════════════════════════════════════════════════════════
    # TAB 1: KONTROL LOTLARI
    # ═══════════════════════════════════════════════════════════════
    tab_lots = ttk.Frame(notebook, padding=10)
    notebook.add(tab_lots, text="  Kontrol Lotlari  ")

    lot_create_frame = ttk.LabelFrame(tab_lots, text="Yangi Lot Yaratish", padding=10)
    lot_create_frame.pack(fill=tk.X, pady=5)

    ttk.Label(lot_create_frame, text="Lot raqami:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=3)
    lot_number_var = tk.StringVar()
    ttk.Entry(lot_create_frame, textvariable=lot_number_var, width=20).grid(row=0, column=1, padx=5, pady=3)

    ttk.Label(lot_create_frame, text="Analizator:").grid(row=0, column=2, sticky=tk.W, padx=5, pady=3)
    analyzer_var = tk.StringVar(value="hematology")
    ttk.Combobox(lot_create_frame, textvariable=analyzer_var, width=18,
                 values=["hematology", "biochemistry"], state="readonly").grid(row=0, column=3, padx=5, pady=3)

    ttk.Label(lot_create_frame, text="Level:").grid(row=0, column=4, sticky=tk.W, padx=5, pady=3)
    level_var = tk.StringVar(value="Level 1")
    ttk.Combobox(lot_create_frame, textvariable=level_var, width=12,
                 values=["Level 1", "Level 2", "Level 3"], state="readonly").grid(row=0, column=5, padx=5, pady=3)

    ttk.Label(lot_create_frame, text="Muddat:").grid(row=0, column=6, sticky=tk.W, padx=5, pady=3)
    expiry_var = tk.StringVar()
    ttk.Entry(lot_create_frame, textvariable=expiry_var, width=12).grid(row=0, column=7, padx=5, pady=3)
    ttk.Label(lot_create_frame, text="(YYYY-MM-DD)", font=("Arial", 8)).grid(row=1, column=7, sticky=tk.W, padx=5)

    ttk.Label(lot_create_frame, text="Izoh:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=3)
    desc_var = tk.StringVar()
    ttk.Entry(lot_create_frame, textvariable=desc_var, width=50).grid(row=1, column=1, columnspan=5, padx=5, pady=3, sticky=tk.W)

    # --- Lotlar ro'yxati ---
    lots_list_frame = ttk.LabelFrame(tab_lots, text="Mavjud Lotlar", padding=5)
    lots_list_frame.pack(fill=tk.X, pady=5)

    lots_cols = ("ID", "Lot", "Analizator", "Level", "Muddat", "Izoh", "Yaratilgan")
    lots_tree = ttk.Treeview(lots_list_frame, columns=lots_cols, show="headings", height=5)
    for c in lots_cols:
        lots_tree.heading(c, text=c)
    lots_tree.column("ID", width=40)
    lots_tree.column("Lot", width=100)
    lots_tree.column("Analizator", width=100)
    lots_tree.column("Level", width=80)
    lots_tree.column("Muddat", width=100)
    lots_tree.column("Izoh", width=250)
    lots_tree.column("Yaratilgan", width=150)
    lots_tree.pack(fill=tk.X)

    def load_lots():
        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM qc_lots ORDER BY created_at DESC")
                return cursor.fetchall()
            finally:
                conn.close()
        def _update(rows):
            for item in lots_tree.get_children():
                lots_tree.delete(item)
            if not rows:
                return
            for r in rows:
                lots_tree.insert('', tk.END, values=(
                    r['id'], r['lot_number'], r['analyzer_type'],
                    r['level'], r.get('expiry_date', ''),
                    r.get('description', '') or '', r['created_at']
                ))
        _run_in_bg(_fetch, _update)

    # --- Parametrlar paneli ---
    params_frame = ttk.LabelFrame(tab_lots, text="Lot Parametrlari (Target va SD kiritish)", padding=5)
    params_frame.pack(fill=tk.BOTH, expand=True, pady=5)

    params_cols = ("ID", "Parametr", "Target (X\u0304)", "1SD", "CV%", "Birlik", "-2SD", "-1SD", "+1SD", "+2SD", "+3SD")
    params_tree = ttk.Treeview(params_frame, columns=params_cols, show="headings", height=12)
    for c in params_cols:
        params_tree.heading(c, text=c)
    params_tree.column("ID", width=40)
    params_tree.column("Parametr", width=120)
    params_tree.column("Target (X\u0304)", width=90)
    params_tree.column("1SD", width=70)
    params_tree.column("CV%", width=60)
    params_tree.column("Birlik", width=70)
    for c in ("-2SD", "-1SD", "+1SD", "+2SD", "+3SD"):
        params_tree.column(c, width=70)

    params_scroll = ttk.Scrollbar(params_frame, orient=tk.VERTICAL, command=params_tree.yview)
    params_tree.configure(yscrollcommand=params_scroll.set)
    params_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    params_scroll.pack(side=tk.RIGHT, fill=tk.Y)

    def load_params(lot_id):
        current_lot_id[0] = lot_id
        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM qc_lot_params WHERE lot_id = %s ORDER BY parameter_name", (lot_id,))
                return cursor.fetchall()
            finally:
                conn.close()
        def _update(rows):
            for item in params_tree.get_children():
                params_tree.delete(item)
            if not rows:
                return
            for r in rows:
                t = float(r['target_value']) if r['target_value'] else 0
                s = float(r['sd_value']) if r['sd_value'] else 0
                cv = round((s / t) * 100, 1) if t > 0 and s > 0 else ''
                m2 = round(t - 2*s, 2) if t and s else ''
                m1 = round(t - s, 2) if t and s else ''
                p1 = round(t + s, 2) if t and s else ''
                p2 = round(t + 2*s, 2) if t and s else ''
                p3 = round(t + 3*s, 2) if t and s else ''
                params_tree.insert('', tk.END, values=(
                    r['id'], r['parameter_name'],
                    t if t else '', s if s else '', cv,
                    r['unit'] or '', m2, m1, p1, p2, p3
                ))
        _run_in_bg(_fetch, _update)

    def on_lot_select(event):
        sel = lots_tree.selection()
        if sel:
            vals = lots_tree.item(sel[0], 'values')
            load_params(vals[0])
    lots_tree.bind("<<TreeviewSelect>>", on_lot_select)

    # --- Target/SD kiritish ---
    param_edit_frame = ttk.Frame(params_frame)
    param_edit_frame.pack(side=tk.BOTTOM, fill=tk.X, pady=5)

    ttk.Label(param_edit_frame, text="Target:").pack(side=tk.LEFT, padx=3)
    target_entry_var = tk.StringVar()
    ttk.Entry(param_edit_frame, textvariable=target_entry_var, width=12).pack(side=tk.LEFT, padx=3)
    ttk.Label(param_edit_frame, text="1SD:").pack(side=tk.LEFT, padx=3)
    sd_entry_var = tk.StringVar()
    ttk.Entry(param_edit_frame, textvariable=sd_entry_var, width=12).pack(side=tk.LEFT, padx=3)

    def save_param_values():
        sel = params_tree.selection()
        if not sel:
            messagebox.showwarning("Diqqat", "Parametrni tanlang!", parent=win)
            return
        param_id = params_tree.item(sel[0], 'values')[0]
        try:
            t = float(target_entry_var.get().strip().replace(',', '.')) if target_entry_var.get().strip() else None
            s = float(sd_entry_var.get().strip().replace(',', '.')) if sd_entry_var.get().strip() else None
        except ValueError:
            messagebox.showwarning("Xato", "Raqam kiriting!", parent=win)
            return
        def _save():
            conn = _db_conn()
            if not conn:
                return
            try:
                cursor = conn.cursor()
                cursor.execute("UPDATE qc_lot_params SET target_value=%s, sd_value=%s WHERE id=%s", (t, s, param_id))
                conn.commit()
            finally:
                conn.close()
        def _done(_):
            if current_lot_id[0]:
                load_params(current_lot_id[0])
        _run_in_bg(_save, _done)

    ttk.Button(param_edit_frame, text="Saqlash", command=save_param_values).pack(side=tk.LEFT, padx=10)

    def on_param_select(event):
        sel = params_tree.selection()
        if sel:
            vals = params_tree.item(sel[0], 'values')
            current_param_id[0] = vals[0]
            target_entry_var.set(str(vals[2]) if vals[2] else '')
            sd_entry_var.set(str(vals[3]) if vals[3] else '')
    params_tree.bind("<<TreeviewSelect>>", on_param_select)

    # Yangi parametr qo'shish
    add_param_frame = ttk.Frame(param_edit_frame)
    add_param_frame.pack(side=tk.RIGHT, padx=20)
    ttk.Label(add_param_frame, text="Yangi:").pack(side=tk.LEFT, padx=3)
    new_param_var = tk.StringVar()
    ttk.Entry(add_param_frame, textvariable=new_param_var, width=15).pack(side=tk.LEFT, padx=3)
    ttk.Label(add_param_frame, text="Birlik:").pack(side=tk.LEFT, padx=3)
    new_unit_var = tk.StringVar()
    ttk.Entry(add_param_frame, textvariable=new_unit_var, width=10).pack(side=tk.LEFT, padx=3)

    def add_param():
        if not current_lot_id[0]:
            messagebox.showwarning("Diqqat", "Avval lotni tanlang!", parent=win)
            return
        pname = new_param_var.get().strip()
        punit = new_unit_var.get().strip()
        if not pname:
            return
        lid = current_lot_id[0]
        def _save():
            conn = _db_conn()
            if not conn:
                return 'no_conn'
            try:
                cursor = conn.cursor()
                cursor.execute("INSERT INTO qc_lot_params (lot_id, parameter_name, unit) VALUES (%s,%s,%s)",
                               (lid, pname, punit))
                conn.commit()
                return 'ok'
            except mysql.connector.IntegrityError:
                return 'duplicate'
            except Exception as e:
                return str(e)
            finally:
                conn.close()
        def _done(res):
            if res == 'duplicate':
                messagebox.showwarning("Xato", f"'{pname}' allaqachon mavjud!", parent=win)
            elif res == 'ok':
                load_params(lid)
                new_param_var.set("")
                new_unit_var.set("")
        _run_in_bg(_save, _done)
    ttk.Button(add_param_frame, text="Qo'shish", command=add_param).pack(side=tk.LEFT, padx=3)

    # Lot yaratish
    def create_lot():
        lot_num = lot_number_var.get().strip()
        a_type = analyzer_var.get()
        lvl = level_var.get()
        exp = expiry_var.get().strip() or None
        desc = desc_var.get().strip() or None
        if not lot_num:
            messagebox.showwarning("Diqqat", "Lot raqamini kiriting!", parent=win)
            return
        def _save():
            conn = _db_conn()
            if not conn:
                return 'no_conn'
            try:
                cursor = conn.cursor()
                cursor.execute("INSERT INTO qc_lots (lot_number, analyzer_type, level, expiry_date, description) VALUES (%s,%s,%s,%s,%s)",
                               (lot_num, a_type, lvl, exp, desc))
                lot_id = cursor.lastrowid
                if a_type == 'hematology':
                    for pn, pu in HEMATOLOGY_PARAMS:
                        cursor.execute("INSERT INTO qc_lot_params (lot_id, parameter_name, unit) VALUES (%s,%s,%s)", (lot_id, pn, pu))
                elif a_type == 'biochemistry':
                    bio_params = _load_bio_params_sync()
                    for pn, pu in bio_params:
                        cursor.execute("INSERT INTO qc_lot_params (lot_id, parameter_name, unit) VALUES (%s,%s,%s)", (lot_id, pn, pu))
                conn.commit()
                return 'ok'
            except mysql.connector.IntegrityError:
                return 'duplicate'
            except Exception as e:
                return str(e)
            finally:
                conn.close()
        def _done(res):
            if res == 'ok':
                messagebox.showinfo("Tayyor", f"Lot '{lot_num}' ({lvl}) yaratildi!", parent=win)
                load_lots()
                lot_number_var.set("")
                expiry_var.set("")
                desc_var.set("")
            elif res == 'duplicate':
                messagebox.showwarning("Xato", f"Lot '{lot_num}' ({a_type}, {lvl}) allaqachon mavjud!", parent=win)
            elif res and res != 'no_conn':
                messagebox.showerror("Xato", res, parent=win)
        _run_in_bg(_save, _done)

    ttk.Button(lot_create_frame, text="Lot Yaratish", command=create_lot).grid(row=1, column=6, padx=5, pady=3)

    def delete_lot():
        sel = lots_tree.selection()
        if not sel:
            return
        vals = lots_tree.item(sel[0], 'values')
        lot_id_del = vals[0]
        if not messagebox.askyesno("Tasdiqlash", f"Lot #{vals[1]} ni o'chirish?", parent=win):
            return
        def _del():
            conn = _db_conn()
            if not conn:
                return
            try:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM qc_lots WHERE id = %s", (lot_id_del,))
                conn.commit()
            finally:
                conn.close()
        _run_in_bg(_del, lambda _: load_lots())

    ttk.Button(lot_create_frame, text="O'chirish", command=delete_lot).grid(row=0, column=8, padx=5, pady=3)

    # ═══════════════════════════════════════════════════════════════
    # TAB 2: NATIJALAR KIRITISH
    # ═══════════════════════════════════════════════════════════════
    tab_results = ttk.Frame(notebook, padding=10)
    notebook.add(tab_results, text="  Natijalar Kiritish  ")

    res_top = ttk.Frame(tab_results)
    res_top.pack(fill=tk.X, pady=5)

    ttk.Label(res_top, text="Lot:").pack(side=tk.LEFT, padx=5)
    res_lot_var = tk.StringVar()
    res_lot_combo = ttk.Combobox(res_top, textvariable=res_lot_var, width=30, state="readonly")
    res_lot_combo.pack(side=tk.LEFT, padx=5)

    ttk.Label(res_top, text="Parametr:").pack(side=tk.LEFT, padx=5)
    res_param_var = tk.StringVar()
    res_param_combo = ttk.Combobox(res_top, textvariable=res_param_var, width=20, state="readonly")
    res_param_combo.pack(side=tk.LEFT, padx=5)

    ttk.Label(res_top, text="Qiymat:").pack(side=tk.LEFT, padx=5)
    res_value_var = tk.StringVar()
    ttk.Entry(res_top, textvariable=res_value_var, width=12).pack(side=tk.LEFT, padx=5)

    ttk.Label(res_top, text="Manba:").pack(side=tk.LEFT, padx=5)
    res_source_var = tk.StringVar(value="manual")
    ttk.Combobox(res_top, textvariable=res_source_var, width=12,
                 values=["manual", "BC-20S", "BK-280"], state="readonly").pack(side=tk.LEFT, padx=5)

    _lot_map = {}
    _param_map = {}

    def refresh_lot_combos():
        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM qc_lots ORDER BY analyzer_type, lot_number")
                return cursor.fetchall()
            finally:
                conn.close()
        def _update(rows):
            _lot_map.clear()
            lot_list = []
            if rows:
                for r in rows:
                    key = f"{r['lot_number']} -- {r['analyzer_type']} ({r['level']})"
                    _lot_map[key] = r['id']
                    lot_list.append(key)
            res_lot_combo['values'] = lot_list
            ov_lot_combo['values'] = lot_list
            rep_lot_combo['values'] = lot_list
        _run_in_bg(_fetch, _update)

    def on_res_lot_changed(event=None):
        lot_key = res_lot_var.get()
        lot_id = _lot_map.get(lot_key)
        if not lot_id:
            return
        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT id, parameter_name, target_value, sd_value FROM qc_lot_params WHERE lot_id=%s ORDER BY parameter_name", (lot_id,))
                return cursor.fetchall()
            finally:
                conn.close()
        def _update(rows):
            _param_map.clear()
            param_list = []
            if rows:
                for r in rows:
                    pn = r['parameter_name']
                    _param_map[pn] = (r['id'], float(r['target_value'] or 0), float(r['sd_value'] or 0))
                    param_list.append(pn)
            res_param_combo['values'] = param_list
        _run_in_bg(_fetch, _update)
    res_lot_combo.bind("<<ComboboxSelected>>", on_res_lot_changed)

    # Info paneli
    res_info_frame = ttk.LabelFrame(tab_results, text="Parametr Ma'lumoti", padding=5)
    res_info_frame.pack(fill=tk.X, pady=5)
    info_labels = {}
    for i, lbl in enumerate(["Target:", "1SD:", "CV%:", "O'rtacha:", "Soni:", "Westgard:"]):
        ttk.Label(res_info_frame, text=lbl, font=("Arial", 9, "bold")).grid(row=0, column=i*2, padx=5)
        v = ttk.Label(res_info_frame, text="--", font=("Arial", 9))
        v.grid(row=0, column=i*2+1, padx=5)
        info_labels[lbl] = v

    # Natijalar jadvali + grafik
    res_table_frame = ttk.Frame(tab_results)
    res_table_frame.pack(fill=tk.BOTH, expand=True, pady=5)

    res_left = ttk.Frame(res_table_frame)
    res_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    res_cols = ("#", "Sana", "Qiymat", "Manba", "Westgard", "Izoh")
    res_tree = ttk.Treeview(res_left, columns=res_cols, show="headings", height=15)
    for c in res_cols:
        res_tree.heading(c, text=c)
    res_tree.column("#", width=40)
    res_tree.column("Sana", width=140)
    res_tree.column("Qiymat", width=90)
    res_tree.column("Manba", width=80)
    res_tree.column("Westgard", width=100)
    res_tree.column("Izoh", width=80)
    res_scroll = ttk.Scrollbar(res_left, orient=tk.VERTICAL, command=res_tree.yview)
    res_tree.configure(yscrollcommand=res_scroll.set)
    res_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    res_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    res_tree.tag_configure('ok', foreground='#008000')
    res_tree.tag_configure('warn', foreground='#FF8C00')
    res_tree.tag_configure('fail', foreground='#FF0000', background='#FFE0E0')

    res_right = ttk.LabelFrame(res_table_frame, text="Levey-Jennings Grafik", padding=5)
    res_right.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
    lj_chart = LJChart(res_right)

    def update_param_info():
        pname = res_param_var.get()
        data = _param_map.get(pname)
        if not data:
            return
        param_id, target, sd = data
        cv = round((sd / target) * 100, 1) if target > 0 and sd > 0 else 0
        info_labels["Target:"].config(text=f"{target}")
        info_labels["1SD:"].config(text=f"{sd}")
        info_labels["CV%:"].config(text=f"{cv}%")

        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (param_id,))
                return cursor.fetchall()
            finally:
                conn.close()

        def _update(rows):
            if rows is None:
                rows = []
            values = [float(r['measured_value']) for r in rows]
            n = len(values)
            mean = round(sum(values) / n, 2) if n > 0 else 0
            info_labels["O'rtacha:"].config(text=f"{mean}")
            info_labels["Soni:"].config(text=f"{n}")
            if n > 0 and target > 0 and sd > 0:
                status, rule = evaluate_westgard(values, target, sd)
                color = "#008000" if status == 'in_control' else "#FF8C00" if status == 'warning' else "#FF0000"
                text = "Nazoratda" if status == 'in_control' else f"! {rule}" if status == 'warning' else f"X {rule}"
                info_labels["Westgard:"].config(text=text, foreground=color)
            else:
                info_labels["Westgard:"].config(text="--", foreground="black")
            # Jadval
            for item in res_tree.get_children():
                res_tree.delete(item)
            for i, r in enumerate(rows, 1):
                ws = r.get('westgard_status', 'in_control')
                ws_d = "OK" if ws == 'in_control' else f"! {r.get('note','')}" if ws == 'warning' else f"X {r.get('note','')}"
                tag = 'ok' if ws == 'in_control' else 'warn' if ws == 'warning' else 'fail'
                dt = r['measured_at'].strftime("%Y-%m-%d %H:%M") if hasattr(r['measured_at'], 'strftime') else str(r['measured_at'])
                res_tree.insert('', tk.END, values=(i, dt, float(r['measured_value']), r.get('source',''), ws_d, r.get('note','') or ''),
                                tags=(str(r['id']), tag))
            # Grafik
            lj_chart.draw(values, target, sd, pname)

        _run_in_bg(_fetch, _update)

    res_param_combo.bind("<<ComboboxSelected>>", lambda e: update_param_info())

    # Natija qo'shish
    def add_result():
        pname = res_param_var.get()
        data = _param_map.get(pname)
        if not data:
            messagebox.showwarning("Diqqat", "Parametrni tanlang!", parent=win)
            return
        param_id, target, sd = data
        try:
            val = float(res_value_var.get().strip().replace(',', '.'))
        except ValueError:
            messagebox.showwarning("Xato", "Qiymat kiriting!", parent=win)
            return
        source = res_source_var.get()

        def _save():
            conn = _db_conn()
            if not conn:
                return
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (param_id,))
                prev = [float(r[0]) for r in cursor.fetchall()]
                prev.append(val)
                status, rule = evaluate_westgard(prev, target, sd) if target > 0 and sd > 0 else ('in_control', '')
                cursor.execute("INSERT INTO qc_results (lot_param_id, measured_value, source, westgard_status, note) VALUES (%s,%s,%s,%s,%s)",
                               (param_id, val, source, status, rule or None))
                conn.commit()
            finally:
                conn.close()

        def _done(_):
            res_value_var.set("")
            update_param_info()
        _run_in_bg(_save, _done)

    ttk.Button(res_top, text="Qo'shish", command=add_result).pack(side=tk.LEFT, padx=10)

    # Analizatordan import
    def import_from_analyzer():
        lot_key = res_lot_var.get()
        lot_id = _lot_map.get(lot_key)
        if not lot_id:
            messagebox.showwarning("Diqqat", "Avval lotni tanlang!", parent=win)
            return

        def _do():
            conn = _db_conn()
            if not conn:
                return None
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT * FROM qc_lots WHERE id=%s", (lot_id,))
                return cursor.fetchone()
            finally:
                conn.close()

        def _then(lot_info):
            if not lot_info:
                return
            if lot_info['analyzer_type'] == 'hematology':
                _import_hematology_qc(lot_id, lot_info['lot_number'], win)
            else:
                _import_biochemistry_qc(lot_id, lot_info['lot_number'], win)

        _run_in_bg(_do, _then)

    ttk.Button(res_top, text="Analizatordan Yuklash", command=import_from_analyzer).pack(side=tk.LEFT, padx=5)

    # Natija o'chirish
    def delete_result():
        sel = res_tree.selection()
        if not sel:
            return
        tags = res_tree.item(sel[0], 'tags')
        if not tags:
            return
        rid = tags[0]
        if not messagebox.askyesno("Tasdiqlash", "Natijani o'chirish?", parent=win):
            return
        def _del():
            conn = _db_conn()
            if not conn:
                return
            try:
                cursor = conn.cursor()
                cursor.execute("DELETE FROM qc_results WHERE id=%s", (rid,))
                conn.commit()
            finally:
                conn.close()
        _run_in_bg(_del, lambda _: update_param_info())

    ttk.Button(res_left, text="Tanlanganni O'chirish", command=delete_result).pack(pady=5)

    # ═══════════════════════════════════════════════════════════════
    # TAB 3: UMUMIY KO'RINISH
    # ═══════════════════════════════════════════════════════════════
    tab_overview = ttk.Frame(notebook, padding=10)
    notebook.add(tab_overview, text="  Umumiy Ko'rinish  ")

    overview_top = ttk.Frame(tab_overview)
    overview_top.pack(fill=tk.X, pady=5)
    ttk.Label(overview_top, text="Lot:").pack(side=tk.LEFT, padx=5)
    ov_lot_var = tk.StringVar()
    ov_lot_combo = ttk.Combobox(overview_top, textvariable=ov_lot_var, width=30, state="readonly")
    ov_lot_combo.pack(side=tk.LEFT, padx=5)
    ttk.Label(overview_top, text="Oy:").pack(side=tk.LEFT, padx=5)
    ov_month_var = tk.StringVar(value=datetime.now().strftime("%Y-%m"))
    ttk.Entry(overview_top, textvariable=ov_month_var, width=10).pack(side=tk.LEFT, padx=5)

    ov_cols = ("Parametr", "Target", "1SD", "CV%", "Soni", "O'rtacha", "Hisob SD", "Hisob CV%", "Oxirgi", "Westgard")
    ov_tree = ttk.Treeview(tab_overview, columns=ov_cols, show="headings", height=20)
    for c in ov_cols:
        ov_tree.heading(c, text=c)
        ov_tree.column(c, width=100)
    ov_tree.column("Parametr", width=110)
    ov_tree.column("Target", width=80)
    ov_tree.column("1SD", width=60)
    ov_tree.column("CV%", width=55)
    ov_tree.column("Soni", width=45)
    ov_scroll = ttk.Scrollbar(tab_overview, orient=tk.VERTICAL, command=ov_tree.yview)
    ov_tree.configure(yscrollcommand=ov_scroll.set)
    ov_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    ov_scroll.pack(side=tk.RIGHT, fill=tk.Y)
    ov_tree.tag_configure('ok', foreground='#008000')
    ov_tree.tag_configure('warn', foreground='#FF8C00')
    ov_tree.tag_configure('fail', foreground='#FF0000', background='#FFE0E0')

    def load_overview():
        lot_key = ov_lot_var.get()
        lot_id = _lot_map.get(lot_key)
        if not lot_id:
            return
        month_str = ov_month_var.get().strip()

        def _fetch():
            conn = _db_conn()
            if not conn:
                return []
            try:
                cursor = conn.cursor(dictionary=True)
                cursor.execute("SELECT id, parameter_name, target_value, sd_value, unit FROM qc_lot_params WHERE lot_id=%s ORDER BY parameter_name", (lot_id,))
                params = cursor.fetchall()
                result = []
                for p in params:
                    pid = p['id']
                    target = float(p['target_value'] or 0)
                    sd = float(p['sd_value'] or 0)
                    if month_str:
                        cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s AND DATE_FORMAT(measured_at,'%%Y-%%m')=%s ORDER BY measured_at", (pid, month_str))
                    else:
                        cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (pid,))
                    values = [float(r['measured_value']) for r in cursor.fetchall()]
                    result.append((p, values))
                return result
            finally:
                conn.close()

        def _update(data):
            for item in ov_tree.get_children():
                ov_tree.delete(item)
            if not data:
                return
            for p, values in data:
                target = float(p['target_value'] or 0)
                sd = float(p['sd_value'] or 0)
                cv = round((sd / target) * 100, 1) if target > 0 and sd > 0 else ''
                n = len(values)
                mean = round(sum(values) / n, 2) if n > 0 else ''
                calc_sd = ''
                calc_cv = ''
                if n >= 2:
                    m = sum(values) / n
                    calc_sd = round(math.sqrt(sum((v - m)**2 for v in values) / (n - 1)), 3)
                    calc_cv = round((calc_sd / m) * 100, 1) if m > 0 else ''
                last_val = round(values[-1], 2) if values else ''
                ws_txt = '--'
                tag = 'ok'
                if values and target > 0 and sd > 0:
                    status, rule = evaluate_westgard(values, target, sd)
                    ws_txt = "OK" if status == 'in_control' else f"! {rule}" if status == 'warning' else f"X {rule}"
                    tag = 'ok' if status == 'in_control' else 'warn' if status == 'warning' else 'fail'
                ov_tree.insert('', tk.END, values=(
                    p['parameter_name'], target or '', sd or '', cv,
                    n, mean, calc_sd, calc_cv, last_val, ws_txt
                ), tags=(tag,))
        _run_in_bg(_fetch, _update)

    ttk.Button(overview_top, text="Yangilash", command=load_overview).pack(side=tk.LEFT, padx=10)

    # ═══════════════════════════════════════════════════════════════
    # TAB 4: HISOBOT
    # ═══════════════════════════════════════════════════════════════
    tab_report = ttk.Frame(notebook, padding=10)
    notebook.add(tab_report, text="  Hisobot / Blanka  ")

    rep_top = ttk.Frame(tab_report)
    rep_top.pack(fill=tk.X, pady=10)
    ttk.Label(rep_top, text="Lot:").pack(side=tk.LEFT, padx=5)
    rep_lot_var = tk.StringVar()
    rep_lot_combo = ttk.Combobox(rep_top, textvariable=rep_lot_var, width=30, state="readonly")
    rep_lot_combo.pack(side=tk.LEFT, padx=5)
    ttk.Label(rep_top, text="Oy:").pack(side=tk.LEFT, padx=5)
    rep_month_var = tk.StringVar(value=datetime.now().strftime("%Y-%m"))
    ttk.Entry(rep_top, textvariable=rep_month_var, width=10).pack(side=tk.LEFT, padx=5)

    def generate_qc_report():
        lot_key = rep_lot_var.get()
        lot_id = _lot_map.get(lot_key)
        if not lot_id:
            messagebox.showwarning("Diqqat", "Lotni tanlang!", parent=win)
            return
        month_str = rep_month_var.get().strip()
        _run_in_bg(lambda: _generate_qc_word_report_sync(lot_id, month_str),
                   lambda path: messagebox.showinfo("Tayyor", f"Hisobot yaratildi:\n{path}", parent=win) if path else None)

    ttk.Button(rep_top, text="Word Blanka Yaratish", command=generate_qc_report).pack(side=tk.LEFT, padx=10)
    ttk.Label(tab_report, text="Lotni va oyni tanlab 'Word Blanka Yaratish' tugmasini bosing.",
              font=("Arial", 11)).pack(pady=30)

    # Tab o'zgarganda sync
    def on_tab_changed(event):
        refresh_lot_combos()
    notebook.bind("<<NotebookTabChanged>>", on_tab_changed)

    # Boshlang'ich yuklash
    load_lots()
    refresh_lot_combos()

    return win


# ══════════════════════════════════════════════════════════════════
# GEMOTOLOGIYA QC IMPORT
# ══════════════════════════════════════════════════════════════════
def _import_hematology_qc(lot_id, lot_number, parent_win):
    import glob as glob_mod
    import re as re_mod
    BC20S_DAT_BASE = r"C:\Program Files (x86)\Mindray\Hematology Analyzer Data Management\LabInfosystem\Dat"
    today_str = datetime.now().strftime("%Y%m")
    dat_folder = os.path.join(BC20S_DAT_BASE, today_str)

    def _do():
        if not os.path.exists(dat_folder):
            return 'no_folder'
        all_files = sorted(glob_mod.glob(os.path.join(dat_folder, "*.txt")), key=os.path.getmtime, reverse=True)
        found_values = {}
        found_file = None
        for fpath in all_files:
            try:
                with open(fpath, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
            except Exception:
                continue
            segments = re_mod.split(r'(?=MSH\|)', content)
            for block in segments:
                lines = block.strip().split('\n')
                sid_found = None
                temp = {}
                for line in lines:
                    line = line.strip('\r\n ')
                    fields = line.split('|')
                    if len(fields) < 2:
                        continue
                    seg = fields[0]
                    if seg == 'OBR' and len(fields) > 3:
                        sid = fields[2].strip() if len(fields) > 2 else ''
                        if not sid and len(fields) > 3:
                            sid = fields[3].strip()
                        if sid == lot_number:
                            sid_found = sid
                    elif seg == 'OBX' and sid_found and len(fields) > 5:
                        tc = fields[3].strip() if len(fields) > 3 else ''
                        tn = ''
                        if '^' in tc:
                            parts = tc.split('^')
                            tc = parts[0]
                            tn = parts[1] if len(parts) > 1 else ''
                        vs = fields[5].strip() if len(fields) > 5 else ''
                        try:
                            if 15000 <= int(tc) <= 15200:
                                continue
                        except ValueError:
                            pass
                        if tn and vs:
                            try:
                                float(vs)
                                temp[tn] = vs
                            except ValueError:
                                pass
                if sid_found and temp:
                    found_values = temp
                    found_file = fpath
                    break
            if found_values:
                break
        if not found_values:
            return 'not_found'
        # Save
        conn = _db_conn()
        if not conn:
            return 'no_conn'
        imported = 0
        try:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT id, parameter_name FROM qc_lot_params WHERE lot_id=%s", (lot_id,))
            pm = {r['parameter_name']: r['id'] for r in cursor.fetchall()}
            # Case insensitive mapping
            pm_upper = {k.upper(): (k, v) for k, v in pm.items()}
            for hl7n, vs in found_values.items():
                key = hl7n.upper()
                if key in pm_upper:
                    pn, pid = pm_upper[key]
                elif key in {k.upper() for k in pm}:
                    pid = pm.get(hl7n, None)
                    if not pid:
                        continue
                else:
                    continue
                val = float(vs)
                cursor.execute("SELECT target_value, sd_value FROM qc_lot_params WHERE id=%s", (pid,))
                pr = cursor.fetchone()
                tgt = float(pr['target_value'] or 0) if pr else 0
                sdd = float(pr['sd_value'] or 0) if pr else 0
                cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (pid,))
                prev = [float(r['measured_value']) for r in cursor.fetchall()]
                prev.append(val)
                st, rl = evaluate_westgard(prev, tgt, sdd) if tgt > 0 and sdd > 0 else ('in_control', '')
                cursor.execute("INSERT INTO qc_results (lot_param_id, measured_value, source, westgard_status, note) VALUES (%s,%s,'BC-20S',%s,%s)",
                               (pid, val, st, rl or None))
                imported += 1
            conn.commit()
            return f"ok:{imported}:{os.path.basename(found_file)}"
        finally:
            conn.close()

    def _done(res):
        if res == 'no_folder':
            messagebox.showinfo("Ma'lumot", f"Papka topilmadi: {dat_folder}", parent=parent_win)
        elif res == 'not_found':
            messagebox.showinfo("Topilmadi", f"Lot '{lot_number}' uchun BC-20S natijasi topilmadi.", parent=parent_win)
        elif res and res.startswith('ok:'):
            parts = res.split(':')
            messagebox.showinfo("Tayyor", f"BC-20S dan {parts[1]} ta QC natija yuklandi.\nFayl: {parts[2]}", parent=parent_win)
    _run_in_bg(_do, _done)


# ══════════════════════════════════════════════════════════════════
# BIOXIMIYA QC IMPORT
# ══════════════════════════════════════════════════════════════════
def _import_biochemistry_qc(lot_id, lot_number, parent_win):
    import glob as glob_mod
    import re as re_mod
    BK280_RAW = r"G:\DASTUR\URIT 50\BK280\RAW_LOGS"

    def _do():
        if not os.path.exists(BK280_RAW):
            return 'no_folder'
        today = datetime.now()
        patterns = [
            os.path.join(BK280_RAW, today.strftime("%Y%m%d"), "*.hl7"),
            os.path.join(BK280_RAW, today.strftime("%Y%m%d"), "*.txt"),
            os.path.join(BK280_RAW, (today - timedelta(days=1)).strftime("%Y%m%d"), "*.hl7"),
            os.path.join(BK280_RAW, (today - timedelta(days=1)).strftime("%Y%m%d"), "*.txt"),
        ]
        all_files = []
        for pat in patterns:
            all_files.extend(glob_mod.glob(pat))
        all_files.sort(key=os.path.getmtime, reverse=True)
        try:
            from biochemistry_window import LIS_CODE_MAP
        except ImportError:
            LIS_CODE_MAP = {}
        found_values = {}
        found_file = None
        for fpath in all_files:
            try:
                with open(fpath, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
            except Exception:
                continue
            segments = re_mod.split(r'(?=MSH\|)', content)
            for block in segments:
                lines = block.strip().split('\n')
                sid_found = None
                temp = {}
                for line in lines:
                    line = line.strip('\r\n ')
                    fields = line.split('|')
                    if len(fields) < 2:
                        continue
                    seg = fields[0]
                    if seg == 'OBR' and len(fields) > 2:
                        sid = fields[2].strip()
                        if sid == lot_number:
                            sid_found = sid
                    elif seg == 'OBX' and sid_found and len(fields) > 5:
                        tc = fields[3].strip() if len(fields) > 3 else ''
                        if '^' in tc:
                            tc = tc.split('^')[0]
                        vs = fields[5].strip() if len(fields) > 5 else ''
                        tn = LIS_CODE_MAP.get(tc, tc)
                        if vs:
                            try:
                                float(vs)
                                temp[tn] = vs
                            except ValueError:
                                pass
                if sid_found and temp:
                    found_values = temp
                    found_file = fpath
                    break
            if found_values:
                break
        if not found_values:
            return 'not_found'
        conn = _db_conn()
        if not conn:
            return 'no_conn'
        imported = 0
        try:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT id, parameter_name FROM qc_lot_params WHERE lot_id=%s", (lot_id,))
            pm = {}
            for r in cursor.fetchall():
                pm[r['parameter_name'].upper()] = (r['id'], r['parameter_name'])
            for tn, vs in found_values.items():
                key = tn.upper()
                if key not in pm:
                    for pk in pm:
                        if key in pk or pk in key:
                            key = pk
                            break
                if key not in pm:
                    continue
                pid = pm[key][0]
                val = float(vs)
                cursor.execute("SELECT target_value, sd_value FROM qc_lot_params WHERE id=%s", (pid,))
                pr = cursor.fetchone()
                tgt = float(pr['target_value'] or 0) if pr else 0
                sdd = float(pr['sd_value'] or 0) if pr else 0
                cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (pid,))
                prev = [float(r['measured_value']) for r in cursor.fetchall()]
                prev.append(val)
                st, rl = evaluate_westgard(prev, tgt, sdd) if tgt > 0 and sdd > 0 else ('in_control', '')
                cursor.execute("INSERT INTO qc_results (lot_param_id, measured_value, source, westgard_status, note) VALUES (%s,%s,'BK-280',%s,%s)",
                               (pid, val, st, rl or None))
                imported += 1
            conn.commit()
            return f"ok:{imported}:{os.path.basename(found_file)}"
        finally:
            conn.close()

    def _done(res):
        if res == 'no_folder':
            messagebox.showinfo("Ma'lumot", f"BK-280 papka topilmadi", parent=parent_win)
        elif res == 'not_found':
            messagebox.showinfo("Topilmadi", f"Lot '{lot_number}' uchun BK-280 natijasi topilmadi.", parent=parent_win)
        elif res and res.startswith('ok:'):
            parts = res.split(':')
            messagebox.showinfo("Tayyor", f"BK-280 dan {parts[1]} ta QC natija yuklandi.\nFayl: {parts[2]}", parent=parent_win)
    _run_in_bg(_do, _done)


# ══════════════════════════════════════════════════════════════════
# WORD HISOBOT (fon thread da)
# ══════════════════════════════════════════════════════════════════
def _generate_qc_word_report_sync(lot_id, month_str):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    conn = _db_conn()
    if not conn:
        return None
    try:
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM qc_lots WHERE id=%s", (lot_id,))
        lot = cursor.fetchone()
        if not lot:
            return None

        cursor.execute("SELECT * FROM qc_lot_params WHERE lot_id=%s ORDER BY parameter_name", (lot_id,))
        params = cursor.fetchall()

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("SIFAT NAZORATI HISOBOTI")
        run.bold = True
        run.font.size = Pt(14)

        info = doc.add_paragraph()
        info.add_run(f"Lot: ").bold = True
        info.add_run(f"{lot['lot_number']}  |  ")
        info.add_run(f"Analizator: ").bold = True
        a_name = "Gemotologik (BC-20S)" if lot['analyzer_type'] == 'hematology' else "Bioximik (BK-280)"
        info.add_run(f"{a_name}  |  ")
        info.add_run(f"Level: ").bold = True
        info.add_run(f"{lot['level']}  |  ")
        info.add_run(f"Oy: ").bold = True
        info.add_run(f"{month_str}  |  ")
        info.add_run(f"Sana: ").bold = True
        info.add_run(datetime.now().strftime('%Y-%m-%d %H:%M'))

        cols = ["#", "Parametr", "Birlik", "Target", "1SD", "CV%", "Soni", "O'rtacha", "Hisob SD", "Westgard"]
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, cn in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = cn
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)

        for idx, p in enumerate(params, 1):
            cursor.execute("SELECT id FROM qc_lot_params WHERE lot_id=%s AND parameter_name=%s", (lot_id, p['parameter_name']))
            pr = cursor.fetchone()
            pid = pr['id'] if pr else None
            target = float(p['target_value'] or 0)
            sd = float(p['sd_value'] or 0)
            cv = round((sd / target) * 100, 1) if target > 0 and sd > 0 else ''
            values = []
            if pid:
                if month_str:
                    cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s AND DATE_FORMAT(measured_at,'%%Y-%%m')=%s ORDER BY measured_at", (pid, month_str))
                else:
                    cursor.execute("SELECT measured_value FROM qc_results WHERE lot_param_id=%s ORDER BY measured_at", (pid,))
                values = [float(r['measured_value']) for r in cursor.fetchall()]
            n = len(values)
            mean = round(sum(values) / n, 2) if n > 0 else ''
            calc_sd = round(math.sqrt(sum((v - sum(values)/n)**2 for v in values) / (n - 1)), 3) if n >= 2 else ''
            ws_txt = ''
            if values and target > 0 and sd > 0:
                status, rule = evaluate_westgard(values, target, sd)
                ws_txt = "Nazoratda" if status == 'in_control' else f"Ogohlantirish: {rule}" if status == 'warning' else f"RAD: {rule}"
            row_cells = table.add_row().cells
            row_data = [str(idx), p['parameter_name'], p['unit'] or '',
                        str(target) if target else '', str(sd) if sd else '', str(cv),
                        str(n), str(mean), str(calc_sd), ws_txt]
            for i, val in enumerate(row_data):
                row_cells[i].text = val
                for para in row_cells[i].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in para.runs:
                        r.font.size = Pt(9)
                        if i == len(cols) - 1:
                            if 'RAD' in val:
                                r.font.color.rgb = RGBColor(255, 0, 0)
                            elif 'Ogohlantirish' in val:
                                r.font.color.rgb = RGBColor(255, 140, 0)
                            else:
                                r.font.color.rgb = RGBColor(0, 128, 0)

        SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
        qc_dir = os.path.join(SCRIPT_DIR, "QC_REPORTS")
        os.makedirs(qc_dir, exist_ok=True)
        filename = f"QC_{lot['lot_number']}_{lot['analyzer_type']}_{month_str}_{datetime.now().strftime('%H%M%S')}.docx"
        docx_path = os.path.join(qc_dir, filename)
        doc.save(docx_path)

        try:
            import importlib
            monoblok = importlib.import_module('monoblok_dastur')
            if hasattr(monoblok, 'save_pdf_to_onedrive'):
                threading.Thread(target=monoblok.save_pdf_to_onedrive, args=(docx_path,), daemon=True).start()
        except Exception:
            pass

        try:
            os.startfile(docx_path)
        except Exception:
            pass

        return docx_path
    except Exception as e:
        print(f"[QC] Hisobot xato: {e}")
        return None
    finally:
        conn.close()


if __name__ == '__main__':
    root = tk.Tk()
    root.withdraw()
    open_qc_window(root)
    root.mainloop()
