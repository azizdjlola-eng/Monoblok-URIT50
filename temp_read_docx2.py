import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

files = [
    ('G:\' + '\u049a' + '\u0438\u043b\u0438\u043d\u0433\u0430\u043d \u0430\u043d\u0430\u043b\u0438\u0437\u043b\u0430\u0440\05.03.2026\260304001324 Ismoilova Iroda.docx'),
    ('G:\' + '\u049a' + '\u0438\u043b\u0438\u043d\u0433\u0430\u043d \u0430\u043d\u0430\u043b\u0438\u0437\u043b\u0430\u0440\04.03.2026\260301001202 Test qilish uchun.docx'),
]

for fpath in files:
    print("\n" + "="*60)
    print(f"FILE: {fpath}")
    print("="*60)
    try:
        doc = Document(fpath)
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"  P{i}: {para.text[:120]}")
        print(f"\n  TABLES: {len(doc.tables)}")
        for ti, table in enumerate(doc.tables):
            print(f"\n  Table {ti+1} ({len(table.rows)} rows x {len(table.columns)} cols):")
            for ri, row in enumerate(table.rows[:10]):
                cells = [c.text[:40] for c in row.cells]
                print(f"    Row{ri}: {cells}")
            if len(table.rows) > 10:
                print(f"    ... ({len(table.rows)-10} more rows)")
    except Exception as e:
        print(f"  ERROR: {e}")
